#!/usr/bin/env python3
"""Audit whether every changed DOCX block/media item is visibly marked in review copy."""
from __future__ import annotations

import hashlib
import json
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn

YELLOW = WD_COLOR_INDEX.YELLOW


def highlighted_paragraph(p) -> bool:
    if any(r.font.highlight_color == YELLOW for r in p.runs):
        return True
    ppr = p._p.pPr
    if ppr is not None:
        shd = ppr.find(qn("w:shd"))
        if shd is not None and (shd.get(qn("w:fill")) or "").upper() in {"FFF2CC", "FFFF00"}:
            return True
    return False


def cell_marked(cell) -> bool:
    if any(highlighted_paragraph(p) for p in cell.paragraphs):
        return True
    tcpr = cell._tc.tcPr
    if tcpr is not None:
        shd = tcpr.find(qn("w:shd"))
        if shd is not None and (shd.get(qn("w:fill")) or "").upper() in {"FFF2CC", "FFFF00"}:
            return True
    return False


def media_hashes(path: Path) -> dict[str, str]:
    with zipfile.ZipFile(path) as zf:
        return {
            name: hashlib.sha256(zf.read(name)).hexdigest()
            for name in zf.namelist()
            if name.startswith("word/media/")
        }


def image_paragraph_media(doc: Document) -> dict[int, list[str]]:
    rels = doc.part.rels
    out: dict[int, list[str]] = {}
    for i, p in enumerate(doc.paragraphs):
        names: list[str] = []
        for blip in p._p.xpath(".//a:blip"):
            rid = blip.get(qn("r:embed"))
            if rid and rid in rels:
                target = str(rels[rid].target_ref).replace("\\", "/")
                names.append("word/" + target.lstrip("/"))
        if names:
            out[i] = names
    return out


def main() -> int:
    if len(sys.argv) != 4:
        raise SystemExit("usage: audit_docx_highlights.py ORIGINAL NORMAL HIGHLIGHT")
    original_path, normal_path, highlight_path = map(Path, sys.argv[1:])
    original, normal, review = map(Document, (original_path, normal_path, highlight_path))

    changed_paragraphs = []
    missing_paragraph_marks = []
    max_p = max(len(original.paragraphs), len(normal.paragraphs))
    for i in range(max_p):
        old = original.paragraphs[i].text if i < len(original.paragraphs) else None
        new = normal.paragraphs[i].text if i < len(normal.paragraphs) else None
        if old != new:
            changed_paragraphs.append(i)
            if i >= len(review.paragraphs) or not highlighted_paragraph(review.paragraphs[i]):
                missing_paragraph_marks.append(i)

    changed_cells = []
    deleted_cells = []
    missing_cell_marks = []
    max_t = max(len(original.tables), len(normal.tables))
    for ti in range(max_t):
        old_t = original.tables[ti] if ti < len(original.tables) else None
        new_t = normal.tables[ti] if ti < len(normal.tables) else None
        review_t = review.tables[ti] if ti < len(review.tables) else None
        old_rows = len(old_t.rows) if old_t else 0
        new_rows = len(new_t.rows) if new_t else 0
        for ri in range(max(old_rows, new_rows)):
            old_cols = len(old_t.rows[ri].cells) if old_t and ri < old_rows else 0
            new_cols = len(new_t.rows[ri].cells) if new_t and ri < new_rows else 0
            for ci in range(max(old_cols, new_cols)):
                old = old_t.cell(ri, ci).text if old_t and ri < old_rows and ci < old_cols else None
                new = new_t.cell(ri, ci).text if new_t and ri < new_rows and ci < new_cols else None
                if old != new:
                    key = [ti, ri, ci]
                    changed_cells.append(key)
                    # A removed row/cell cannot be highlighted in a clean review
                    # copy because it no longer exists. Record it explicitly as
                    # a structural deletion; only surviving/new cells require a
                    # visible yellow mark.
                    if new is None:
                        deleted_cells.append(key)
                        continue
                    marked = bool(
                        review_t
                        and ri < len(review_t.rows)
                        and ci < len(review_t.rows[ri].cells)
                        and cell_marked(review_t.cell(ri, ci))
                    )
                    if not marked:
                        missing_cell_marks.append(key)

    old_media, new_media = media_hashes(original_path), media_hashes(normal_path)
    changed_media = sorted(
        name for name in set(old_media) | set(new_media) if old_media.get(name) != new_media.get(name)
    )
    image_map = image_paragraph_media(normal)
    media_paragraphs = {
        name: i for i, names in image_map.items() for name in names if name in changed_media
    }
    missing_media_marks = []
    for name in changed_media:
        i = media_paragraphs.get(name)
        # A changed image is visibly marked when its drawing paragraph is shaded/highlighted.
        if i is None or i >= len(review.paragraphs) or not highlighted_paragraph(review.paragraphs[i]):
            missing_media_marks.append({"media": name, "paragraph": i})

    result = {
        "changed_paragraphs": changed_paragraphs,
        "missing_paragraph_marks": missing_paragraph_marks,
        "changed_cells": changed_cells,
        "deleted_cells": deleted_cells,
        "missing_cell_marks": missing_cell_marks,
        "changed_media": changed_media,
        "changed_media_paragraphs": media_paragraphs,
        "missing_media_marks": missing_media_marks,
        "complete": not (missing_paragraph_marks or missing_cell_marks or missing_media_marks),
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result["complete"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
