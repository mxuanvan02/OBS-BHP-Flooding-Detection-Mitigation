from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import hashlib, zipfile, os, json, csv, sys

base=Path(__file__).resolve().parents[1]
src=base/'deliverables/LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
original=base/'deliverables/LuanVan_ThS_NguyenQuangTin_BAN_GOC_01072026.docx'
original_sha256='a5cb463bd902422cee6e3e243157b238de02aedab4b881c957cd0b650480637e'
# The stable updated source is used to rebuild content; the immutable original is
# the sole comparison baseline for the review/highlight copy.
assert original.is_file(), f'missing original DOCX: {original}'
assert hashlib.sha256(original.read_bytes()).hexdigest() == original_sha256, 'original DOCX hash mismatch'
out=base/'LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
high=base/'LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726_HIGHLIGHT.docx'

def setp(p,text):
    if p.runs:
        p.runs[0].text=text
        for r in p.runs[1:]: r.text=''
    else: p.add_run(text)

doc=Document(src)
# Paragraph indices verified against the actual source DOCX.
repl = {14: 'Bảng 3.1. Kết quả tái phân tích bốn mô hình cơ sở trên UCI404 bằng StratifiedGroupKFold (25 lượt đánh giá/mô hình).\t46', 15: 'Bảng 3.2. Bốn kịch bản mô phỏng đối chứng.\t49', 16: 'Bảng 3.3. Tác động của tấn công lên các chỉ số mạng (tám hạt giống).\t50', 17: 'Bảng 3.4. Trạng thái kiểm toán bộ chuẩn đánh giá theo cửa sổ từ dấu vết nguyên bản và tiêu chí không suy biến.\t53', 18: 'Bảng 3.5. Trạng thái bằng chứng cho đường cong phát hiện theo cường độ tấn công.\t53', 19: 'Bảng 3.6. Thông lượng TCP hợp pháp theo bốn kịch bản direct-BHP (khoảng mô tả trên tám seed).\t55', 34: 'Hình 3.1. Kết quả phép thử từng đặc trưng trên UCI404 dưới giao thức nhóm bản sao chính xác.\t46', 35: 'Hình 3.2. Độ quan trọng hoán vị ngoài mẫu của các đặc trưng UCI404, dùng macro-F1 làm thước đo.\t47', 36: 'Hình 3.3. Kiến trúc mục tiêu phát hiện–quyết định–ứng phó tại nút biên; ma trận hiện tại chỉ kiểm chứng cơ chế kiểm soát BHP trực tiếp trước khi đặt trước tài nguyên và hai cấu hình ứng phó.\t49', 37: 'Hình 3.4. Tác động của BHP điều khiển trực tiếp không kèm chùm dữ liệu lên thông lượng TCP hợp pháp giữa S0 và S1 trên tám hạt giống ngẫu nhiên cố định.\t51', 38: 'Hình 3.5. Số gói TCP hợp pháp theo từng hạt giống ngẫu nhiên trong bốn kịch bản; mỗi đường biểu diễn tám lượt chạy từ môi trường NS-2.35+nOBS nguyên bản.\t52', 39: 'Hình 3.6. Phân bố các quyết định của cơ chế kiểm soát BHP trực tiếp theo hai cấu hình ứng phó, tổng hợp trên tám hạt giống ngẫu nhiên.\t54', 40: 'Hình 3.7. Hiệu quả của hai cấu hình kiểm soát BHP trực tiếp trên tám hạt giống ngẫu nhiên cố định.\t56', 322: 'Đối với bài toán phát hiện tấn công ngập lụt BHP, dữ liệu chứa nhiều bản sao hoặc các biến liên quan chặt với cơ chế tạo nhãn có thể tạo ra đánh giá lạc quan. Vì vậy, kết quả gần tuyệt đối của một mô hình đơn giản được kiểm tra bằng giao thức nhóm bản sao, phép thử từng đặc trưng và độ quan trọng hoán vị ngoài mẫu [27].', 325: 'Khoảng trống thứ nhất liên quan đến tính toàn vẹn của dữ liệu đánh giá. Các kết quả gần tuyệt đối trên UCI404 cần được kiểm tra dưới một giao thức hạn chế sự giao nhau của các bản sao chính xác giữa tập huấn luyện và tập kiểm tra. Luận văn vì vậy thực hiện tái phân tích có kiểm soát theo nhóm bản sao, kết hợp phép thử từng đặc trưng và độ quan trọng hoán vị ngoài mẫu.', 326: 'Khoảng trống thứ hai liên quan đến dữ liệu đánh giá phản ánh điều kiện vận hành. Dấu vết thực nghiệm từ môi trường NS-2.35+nOBS nguyên bản được tổ chức thành bộ chuẩn đánh giá theo cửa sổ với tiêu chí kiểm định tính không suy biến, nhãn và đặc trưng khả dụng theo quan hệ nhân quả, đồng thời tách biệt theo từng lượt chạy. Các kết quả và tiêu chí tương ứng được trình bày trong Chương 3.', 327: 'Khoảng trống thứ ba là sự kết nối giữa phát hiện và ứng phó. Kiến trúc mục tiêu nối dữ liệu giám sát, quyết định và hành động tại nút biên. Luận văn hiện thực và đánh giá cơ chế kiểm soát xác định cùng hai cấu hình ứng phó; mô hình học máy trực tuyến, cơ chế phản hồi và phục hồi đầy đủ được xác định thành hướng mở rộng tiếp theo.', 328: 'Ba khoảng trống trên được dùng làm tiêu chí kiểm toán Chương 3. Mỗi kết luận chỉ được giữ ở mức mà dữ liệu, mã nguồn và dấu vết thực nghiệm hiện có trực tiếp hỗ trợ.', 332: 'Chương này đối chiếu các luận điểm thực nghiệm với những bằng chứng có thể tái kiểm. Quá trình được tổ chức thành hai pha độc lập về nguồn dữ liệu: tái phân tích UCI404 theo quy trình hạn chế rò rỉ do bản sao và mô phỏng bằng môi trường NS-2.35+nOBS nguyên bản đối với BHP điều khiển trực tiếp không kèm chùm dữ liệu.', 333: 'Pha thứ nhất tái phân tích bốn mô hình học máy có mã hiện thực đầy đủ gồm cây quyết định, SVM-RBF, KNN và Gaussian Naïve Bayes. Nhánh PSO–SVM được giữ ở phần tổng quan đối chiếu; pha tái phân tích tập trung vào các mô hình có mã tối ưu, cấu hình tìm kiếm và đầu ra dự đoán có thể kiểm tra. Mục tiêu của pha này là đánh giá mức độ phụ thuộc của kết quả vào bản sao và các đặc trưng có nguy cơ hậu nghiệm.', 334: 'Pha thứ hai chuyển sang mô phỏng động mạng OBS nhằm đánh giá tác động của nhánh BHP điều khiển trực tiếp tại nút biên và các chế độ ứng phó tương ứng. Ma trận mới sử dụng môi trường NS-2.35 kết hợp nOBS, gồm bốn kịch bản S0, S1, S2-rate-limit và S2-cách ly trên tám hạt giống ngẫu nhiên cố định. Nhánh này tạo BHP điều khiển không kèm chùm dữ liệu và thực hiện kiểm soát chấp nhận trước khi đặt trước tài nguyên; đây là thiết kế thực nghiệm riêng cho đường xử lý BHP trực tiếp, song song với nhánh UDP/CBR tạo chùm dữ liệu hợp lệ trước đây.', 335: 'Nguyên tắc phân tích là tách riêng số liệu báo cáo trong bản cũ, kết quả tái phân tích UCI404 và kết quả BHP điều khiển trực tiếp từ môi trường NS-2.35+nOBS. Một kết luận chỉ được sử dụng khi có nguồn vào, giao thức, mã phân tích và đầu ra kiểm tra được; những mục tiêu chưa đủ bằng chứng được trình bày như giới hạn thay vì suy diễn từ bằng chứng thực nghiệm khác.', 337: 'Bộ dữ liệu ở pha thứ nhất là OBS-Network-DataSet trên UCI, gồm 1.075 mẫu, 21 biến dự báo và nhãn Class bốn lớp. Dữ liệu có 860 hàng lặp chính xác, chỉ 215 vector dự báo duy nhất và 15 giá trị thiếu tại Packet_lost; Node Status là biến danh mục, các biến còn lại được xử lý như biến số. Phân bố lớp lần lượt là 500, 300, 155 và 120 mẫu.', 338: 'Do dữ liệu vừa mất cân bằng vừa có mức lặp cao, đánh giá báo cáo đồng thời độ chính xác, độ chính xác cân bằng, macro-F1 và MCC. Các bản sao chính xác được gom nhóm trước khi chia phần để cùng một vector dự báo không xuất hiện ở cả tập huấn luyện và tập kiểm tra.', 339: 'Một số biến như Flood Status, Node Status và các đại lượng tổng hợp qua nhiều lần chạy có thể chứa thông tin hậu nghiệm hoặc phản ánh trực tiếp cơ chế tạo nhãn. Phân tích dưới đây kiểm tra nguy cơ proxy/target-policy leakage và xác định phạm vi diễn giải của kết quả theo từng giao thức đánh giá.', 342: 'Bốn mô hình cơ sở có mã và đầu ra đầy đủ gồm DecisionTree, SVM-RBF, KNN và GaussianNB. Quy trình đánh giá sử dụng StratifiedGroupKFold với năm phần chia và năm hạt giống ngẫu nhiên 17, 42, 73, 101 và 2026, tạo 25 lượt đánh giá cho mỗi mô hình. Nhóm được xác định bằng hàm băm của toàn bộ 21 biến dự báo để bảo đảm các bản sao chính xác được phân tách giữa tập huấn luyện và tập kiểm tra của một phần chia. Các bước điền giá trị thiếu, mã hóa một-trong-K cho Node Status và chuẩn hóa đều được ước lượng riêng trên từng tập huấn luyện. Phần tái phân tích tập trung vào các mô hình cơ sở; PSO–SVM và phép đo độ trễ được dành cho hướng mở rộng có artifact tương ứng.', 343: '3.3.2 Kết quả tái phân tích bốn mô hình', 344: 'Bảng 3.1 trình bày kết quả trung bình của 25 lượt đánh giá theo giao thức có kiểm soát nhóm bản sao. DecisionTree đạt macro-F1 0,8082 và MCC 0,6239; SVM-RBF lần lượt đạt 0,7680 và 0,6095. Không mô hình nào đạt kết quả hoàn hảo. Sự khác biệt với các con số cũ cho thấy kết luận phụ thuộc đáng kể vào cách xử lý bản sao và giao thức phân chia dữ liệu.', 345: 'Bảng 3.1. Kết quả tái phân tích bốn mô hình cơ sở trên UCI404 bằng StratifiedGroupKFold (25 lượt đánh giá/mô hình).', 347: 'Phân tích chất lượng dữ liệu gồm ba lớp: thống kê bản sao và giá trị thiếu; đánh giá cây quyết định với từng đặc trưng riêng lẻ dưới cùng giao thức nhóm; và độ quan trọng hoán vị ngoài mẫu của rừng ngẫu nhiên. Các lớp kiểm tra này nhằm xác định mức độ phụ thuộc và nguy cơ học theo proxy, không mặc định coi mọi điểm số cao là rò rỉ nhãn tất định.', 348: 'Trong phép thử từng đặc trưng, Flood Status là biến mạnh nhất nhưng chỉ đạt độ chính xác trung bình 0,7106, macro-F1 0,7845 và MCC 0,5692. Không có bằng chứng tái lập cho tuyên bố cũ về nhiều đặc trưng riêng lẻ đạt gần hoàn hảo. Kết quả đầy đủ được thể hiện trên Hình 3.1.', 350: 'Hình 3.1. Kết quả phép thử từng đặc trưng trên UCI404 dưới giao thức nhóm bản sao chính xác.', 351: 'Độ quan trọng hoán vị được tính ngoài mẫu với rừng ngẫu nhiên 200 cây. Flood Status có mức giảm macro-F1 trung bình 0,2177 khi bị hoán vị; các biến còn lại gần 0 hoặc âm. Kết quả này cho thấy mô hình phụ thuộc mạnh vào Flood Status trong quy trình đánh giá đã công bố, đồng thời xác định biến này là trọng tâm cần kiểm tra về quan hệ với nhãn. Hình 3.2 trình bày toàn bộ phân bố độ quan trọng.', 353: 'Hình 3.2. Độ quan trọng hoán vị ngoài mẫu của các đặc trưng UCI404, dùng macro-F1 làm thước đo.', 354: 'Tổng hợp ba lớp kiểm tra cho thấy UCI404 có rủi ro nghiêm trọng về phụ thuộc do bản sao, biến hậu nghiệm và target-policy/proxy leakage. Giao thức tái phân tích cung cấp kết quả kiểm toán dữ liệu ngoại tuyến, với các chỉ số mô hình và đặc trưng được trình bày theo phạm vi giao thức nhóm bản sao. Bộ dữ liệu này tạo cơ sở cho các bước đánh giá mở rộng về năng lực phát hiện và khả năng triển khai trong mạng OBS.', 356: 'Kết quả UCI404 cho phép so sánh mô hình cơ sở trong phạm vi ngoại tuyến. Quyết định triển khai cần đồng thời dựa trên bộ chuẩn đánh giá với đặc trưng khả dụng trực tuyến, độ trễ từ thu nhận đặc trưng đến thực thi hành động, kích thước mô hình, tài nguyên và chi phí cảnh báo sai; đây là các tiêu chí của giai đoạn đánh giá vận hành tiếp theo.', 357: 'DecisionTree có macro-F1 cao nhất trong bốn mô hình cơ sở đã tái phân tích, song chênh lệch này chưa đủ làm căn cứ lựa chọn mô hình cho triển khai vận hành. Chưa có phép đo độ trễ p50/p95/p99, kích thước mô hình hoặc chuỗi từ bộ phát hiện đến cơ cấu chấp hành trong cùng quy trình xử lý; vì vậy luận văn không tuyên bố đã lựa chọn được mô hình triển khai.', 358: '3.5 Kiến trúc mục tiêu và cơ chế kiểm soát đã hiện thực', 359: '3.5.1 Ranh giới giữa kiến trúc mục tiêu và bằng chứng thực nghiệm', 360: 'Về kiến trúc, quyết định ứng phó phù hợp hơn khi đặt tại nút biên, nơi gói điều khiển vẫn được xử lý điện tử. Vòng khép kín mục tiêu gồm dữ liệu giám sát, phát hiện, quyết định, cơ cấu chấp hành và phản hồi. Tuy nhiên, bằng chứng thực nghiệm hiện tại chỉ hiện thực bộ sinh BHP trực tiếp, cơ chế kiểm soát xác định trước khi đặt trước tài nguyên và hai cấu hình ứng phó; chưa tích hợp mô hình học máy trực tuyến hay cơ chế phản hồi và phục hồi đầy đủ.', 361: 'Cơ chế kiểm soát quan sát các trường đồng thời của gói điều khiển trực tiếp, kiểm tra tính hợp lệ và áp dụng ngân sách sự kiện cùng chi phí đặt trước tài nguyên trước khi gọi thủ tục đặt trước tài nguyên. Kết quả quyết định dẫn đến ALLOW, DROP_OVER_PROFILE, QUARANTINE hoặc RELEASE tùy cấu hình ứng phó. Chuỗi kiểm toán xác nhận thứ tự tạo–quan sát–quyết định–tác động–kết quả; đây là đường xử lý điều khiển dựa trên trạng thái và ngân sách token chứ không phải đầu ra của cây quyết định hay PSO–SVM.', 362: 'Cơ chế đã chạy không được đồng nhất với RFC 2698 đầy đủ: bằng chứng hiện có chưa có cặp CIR/PIR, CBS/PBS và ngữ nghĩa phân màu hoàn chỉnh; cũng chưa có danh sách xám, cơ chế lùi thời gian theo hàm mũ, cơ chế trễ hai ngưỡng và bộ điều khiển phục hồi. Các thành phần này thuộc kiến trúc mục tiêu và là phần cần phát triển tiếp.', 364: 'Hình 3.3. Kiến trúc mục tiêu phát hiện–quyết định–ứng phó tại nút biên; ma trận hiện tại chỉ kiểm chứng cơ chế kiểm soát BHP trực tiếp trước khi đặt trước tài nguyên và hai cấu hình ứng phó.', 365: '3.5.2 Mô hình mạng và bốn kịch bản đối chứng', 367: 'Để đánh giá có đối chứng, ma trận gồm bốn ô kịch bản: S0 là trạng thái nền không tấn công; S1 có ngập lụt BHP trực tiếp với cơ chế kiểm soát đối chứng không giới hạn thực tế (ngân sách 1e9); S2-rate-limit bật giới hạn theo ngân sách BHP; và S2-cách ly bật cách ly tại nút biên. Hai biến thể S2 dùng cùng tải, hạt giống ngẫu nhiên và cấu hình mạng như S1 để bảo đảm so sánh theo cặp. Cấu trúc bốn ô được tóm tắt trong Bảng 3.2.', 370: 'Công cụ mô phỏng là NS-2 phiên bản 2.35 kết hợp mô-đun nOBS trên môi trường Linux. Ma trận chính gồm 32 lượt chạy: bốn kịch bản nhân với tám hạt giống ngẫu nhiên cố định 101, 202, 303, 404, 505, 606, 707 và 808; mỗi lượt kéo dài 5 giây. Tô-pô tái dựng gồm bảy nút, liên kết quang 1000 Mb/s, hai luồng TCP hợp pháp với đường truy nhập 155 Mb/s và tám nguồn BHP trực tiếp. Mỗi nguồn phát BHP điều khiển không kèm chùm dữ liệu, với kích thước khai báo 1000 byte và một gói tin; tốc độ danh nghĩa là 12 Mb/s/nguồn, còn tốc độ hiệu dụng theo tám seed nằm trong khoảng 9,8926–14,2687 Mb/s/nguồn, trung bình 12,3153 Mb/s/nguồn. Cơ chế kiểm soát dùng ngân sách số sự kiện và ngân sách chi phí đặt trước tài nguyên, không phải CIR payload 4 Mb/s. Các khoảng 95% trong phần này chỉ là khoảng mô tả trên tám hạt giống ngẫu nhiên cố định, không phải ước lượng bất định tổng quát.', 372: 'Trong ma trận BHP trực tiếp, S1 làm số gói TCP hợp pháp nhận được giảm từ trung bình 48.678 ở S0 xuống 15.034,25 ở S1, tương ứng giảm 69,12%; số byte TCP hợp pháp giảm từ 50.623.120 xuống 15.633.620, tương ứng giảm 69,12%. Số cặp chùm quang giảm từ 4.917 xuống 2.060,5, tương ứng giảm 58,09%; số lượt đặt trước tài nguyên thành công giảm từ 19.606 xuống 8.216,5, tương ứng giảm 58,09%. Số lượt loại bỏ dữ liệu tường minh bằng 0 ở S0 và S1. Kết quả phù hợp với cơ chế trong đó các BHP trực tiếp, không kèm chùm dữ liệu, cạnh tranh tài nguyên trên đường xử lý đặt trước; không được diễn giải là tải UDP tạo thêm chùm dữ liệu hợp lệ.', 377: 'Nhánh khảo sát quét theo cường độ để đánh giá phát hiện và tác động trước đây không được dùng để suy luận cho ma trận BHP trực tiếp mới. Ma trận được báo cáo ở đây là một thiết kế ghép cặp với hạt giống ngẫu nhiên cố định gồm bốn ô, nhằm cô lập tác động của việc chấp nhận gói điều khiển trực tiếp trước khi đặt trước tài nguyên và hiệu quả của hai cấu hình kiểm soát. Vì vậy, các đường cong 5–50 Mb/s, điểm đảo chiều và số liệu từ phép quét cường độ không có cùng nguồn gốc truy xuất với ma trận 32 ô không được giữ lại.', 379: 'Hình 3.5. Số gói TCP hợp pháp theo từng hạt giống ngẫu nhiên trong bốn kịch bản; mỗi đường biểu diễn tám lượt chạy từ môi trường NS-2.35+nOBS nguyên bản.', 380: '3.6.3 Kiểm tra bộ chuẩn đánh giá phát hiện theo cửa sổ', 381: 'Ma trận direct-BHP không được dùng để tuyên bố đã xây dựng bộ chuẩn đánh giá học máy không suy biến. Các kết quả ML UCI và bộ chuẩn đánh giá nguyên bản theo cửa sổ được giữ ở nhánh audit riêng; trong đó PSO-SVM, độ trễ triển khai và bộ phát hiện trực tuyến chưa có artifact đầy đủ. Ma trận BHP trực tiếp chỉ kiểm chứng chuỗi xử lý nhân quả của tạo–quan sát–quyết định–tác động–kết quả của gói điều khiển và đáp ứng của mạng.', 382: 'Các phép thử direct-BHP kiểm tra rằng control được tạo, quan sát, quyết định và tác động theo đúng thứ tự thời gian; hành động DROP hoặc QUARANTINE xảy ra trước lần gọi thủ tục đặt trước tài nguyên tương ứng. Đây là tiêu chí kiểm định nhân quả của guard, không phải phép đánh giá độ chính xác của một bộ phân loại.', 383: 'Nhánh BHP trực tiếp dùng bộ điều khiển trạng thái/ngân sách token xác định với dữ liệu quan sát đồng thời. Cơ chế kiểm soát không được gọi là bộ phát hiện ML vận hành; các tuyên bố về accuracy, MCC, latency hoặc lựa chọn cây quyết định chỉ được dùng nếu có bằng chứng thực nghiệm và giao thức độc lập tương ứng.', 384: 'Do đó, ma trận BHP trực tiếp được xem là thực nghiệm kiểm chứng kiểm soát chấp nhận và giảm thiểu ở mức mạng, không phải benchmark huấn luyện bộ phát hiện. Các kết luận chỉ giới hạn trong tô-pô bảy nút, cấu hình lưu lượng đã khai báo, thời lượng 5 giây và tám hạt giống ngẫu nhiên cố định.', 385: 'Bảng 3.4. Trạng thái kiểm toán bộ chuẩn đánh giá theo cửa sổ từ dấu vết nguyên bản và tiêu chí không suy biến.', 386: '3.6.4 Khoảng trống đánh giá khả năng phát hiện theo cường độ', 387: 'Ma trận direct-BHP không cung cấp đường cong khả năng phát hiện theo độ ẩn. Các số liệu MCC theo mức 1–35 Mb/s của bản cũ không có dấu vết thô và chuỗi nguồn gốc dữ liệu khớp với thí nghiệm mới, nên không được trình bày như kết quả tái lập.', 389: 'Vì chưa có benchmark bộ phát hiện trực tuyến không suy biến và chưa có phép đo độ trễ đầu cuối, luận văn không suy diễn từ ma trận này rằng một mô hình mạnh hơn là cần thiết ở vùng tấn công ẩn. Đây là giới hạn bằng chứng và là hướng thực nghiệm tiếp theo.', 388: 'Bảng 3.5. Trạng thái bằng chứng cho đường cong phát hiện theo cường độ tấn công.', 391: 'Hình 3.6. Phân bố các quyết định của cơ chế kiểm soát BHP trực tiếp theo hai cấu hình ứng phó, tổng hợp trên tám hạt giống ngẫu nhiên.', 392: '3.6.5 Hiệu quả của hai chế độ kiểm soát direct-BHP tại nút biên', 393: 'Hai chế độ được hiện thực trong môi trường NS-2.35+nOBS nguyên bản tại nút biên. Cơ chế kiểm soát quan sát trực tiếp các gói điều khiển BHP, áp dụng ngân sách số sự kiện và ngân sách chi phí đặt trước tài nguyên trước lần đặt trước tài nguyên đầu tiên. S2-rate-limit cho phép một phần nhỏ gói điều khiển theo cấu hình rồi chuyển các gói điều khiển vượt cấu hình vào cách ly; S2-cách ly dùng profile cách ly chặt hơn. Đây là đáp ứng của một state/ngân sách token guard với dữ liệu quan sát đồng thời, không phải việc nhúng trực tiếp cây quyết định hay mô hình PSO-SVM vào vòng chạy NS-2.', 394: 'Trong ma trận BHP trực tiếp, S0 chỉ có hai luồng TCP hợp pháp; S1 bổ sung tám nguồn BHP trực tiếp không kèm chùm dữ liệu và dùng cơ chế kiểm soát đối chứng với ngân sách 1e9 nên không giới hạn thực tế; hai S2 giữ nguyên traffic và seed của S1 nhưng bật cấu hình kiểm soát tương ứng. Direct BHP dùng dấu mốc vòng đời nội bộ để đi vào đường xử lý điều khiển tổng hợp; marker này không phải nhãn quan sát của bộ phát hiện. Do đó, kết quả đo hiệu quả giảm thiểu không chứng minh tỷ lệ dương tính giả đối với các gói điều khiển nOBS hợp pháp, vì các control hợp pháp không đi vào đường xử lý kiểm soát BHP trực tiếp tường minh.', 395: 'Bảng 3.6 trình bày số gói TCP hợp pháp theo bốn ô kịch bản. Cả 32 lượt chạy đều kết thúc thành công; mỗi ô có tám seed. Dấu vết thô out.tr được phân tích độc lập; chuỗi kiểm toán trong bhp_audit.log được bộ kiểm tra kiểm tra; kết quả được đối chiếu với validation.json và validation.rerun.json. Các khoảng trong bảng là khoảng mô tả trên tập hạt giống ngẫu nhiên cố định.', 399: 'S2-rate-limit nâng số gói TCP hợp pháp từ 15.034,25 ở S1 lên 48.678, đạt 100% mức S0 và khôi phục toàn bộ phần TCP bị mất trong cấu hình khảo sát. S2-cách ly cũng đạt 48.678 gói, tương đương 100% mức S0 và khôi phục 100% phần bị mất. Hai chế độ có cùng kết quả vận hành mạng trong cả tám seed; không có cơ sở kết luận chế độ nào ưu việt hơn về thông lượng.', 400: 'Khác biệt giữa hai chế độ nằm ở dấu vết hành động chứ không nằm ở số byte TCP cuối cùng. Giới hạn tốc độ ghi nhận 48 gói điều khiển được chấp nhận, 144 DROP_OVER_PROFILE, 433.328 QUARANTINE và 32 RELEASE trên toàn ma trận; cách ly ghi nhận 16 admitted, 144 DROP_OVER_PROFILE và 433.360 QUARANTINE. Hai S2 đều làm giảm mạnh số BHP trực tiếp được chấp nhận so với S1, trong khi hai luồng TCP hợp pháp không thấp hơn S1 và trở lại mức S0.', 401: 'Tổng hợp lại, ma trận 32 lượt chạy cho thấy trong cấu hình direct-BHP đã khai báo, cả rate-limit và cách ly đều ngăn phần lớn BHP điều khiển không kèm chùm dữ liệu trước đặt trước tài nguyên và phục hồi thông lượng TCP hợp pháp về đúng mức S0. Kết luận chỉ áp dụng cho tô-pô bảy nút, cấu hình lưu lượng, thời lượng 5 giây và tám hạt giống ngẫu nhiên cố định; không ghi nhận ảnh hưởng ngoài mục tiêu ở đây chỉ là kết quả của hai flow được quan sát và phạm vi guard đã khai báo, không phải tỷ lệ dương tính giả tổng quát.', 404: 'Nhận định thứ nhất là sự phân biệt giữa các nhánh bằng chứng. UCI404 là kiểm toán ngoại tuyến của một bộ chuẩn đánh giá có nguy cơ trùng lặp, phụ thuộc dữ liệu và rò rỉ do cơ chế tạo nhãn hoặc biến đại diện; ma trận BHP trực tiếp nguyên bản là thực nghiệm admission/giảm thiểu ở mức mạng. Không được dùng kết quả của một nhánh để chứng minh kết luận của nhánh kia.', 405: 'Nhận định thứ hai là bản chất đóng góp của luận văn. Đóng góp thực nghiệm hiện được hỗ trợ gồm: kiểm toán rò rỉ và tính tái lập của UCI404 theo giao thức có kiểm soát rò rỉ; xây dựng native đường xử lý BHP trực tiếp có kiểm toán nhân quả; và đánh giá hai chế độ ứng phó tại nút biên bằng ma trận 32 ô. Ma trận mới cho thấy S1 giảm 69,12% byte TCP hợp pháp, còn cả hai S2 phục hồi 100% mức S0 trong cấu hình khảo sát.', 406: 'Nhận định thứ ba là phạm vi và giới hạn suy luận. Direct BHP producer dùng dấu mốc vòng đời để chọn đường xử lý điều khiển tổng hợp; vì vậy kết quả không chứng minh bộ phát hiện inference từ wire-visible evidence, không phải tái lập nguyên trạng của PSO-SVM hay closed-loop ML deployment. Audit nguồn gốc truy xuất, tính nhất quán với nguồn sinh dữ liệu, legitimate-control khả năng đồng tồn tại và generalization trên nhiều tô-pô vẫn cần được mở rộng.', 409: 'Luận văn đạt được các kết quả thực nghiệm được kiểm chứng trong phạm vi bằng chứng và đầu ra hiện có; các mục tiêu chưa có đủ bằng chứng được nêu rõ là giới hạn, không được nâng thành kết luận đã hoàn tất.', 410: 'Ba kết quả được hỗ trợ trực tiếp là: tái phân tích UCI404 theo giao thức nhóm bản sao và chỉ ra rủi ro phụ thuộc dữ liệu/biến đại diện; xây dựng gói điều khiển trực tiếp-only BHP path trong môi trường NS-2.35+nOBS nguyên bản với audit thứ tự trước đặt trước tài nguyên; và đánh giá hai cấu hình kiểm soát bằng ma trận 32 lượt chạy. Bộ chuẩn đánh giá học máy nguyên bản không suy biến, lựa chọn mô hình vận hành và closed-loop bộ phát hiện-driven chưa đạt tiêu chí kiểm định bằng chứng.', 411: 'Thứ nhất, ma trận BHP trực tiếp nguyên bản trên NS-2.35+nOBS gồm 32 lượt chạy hợp lệ. Trên tám hạt giống ngẫu nhiên cố định, số gói TCP hợp pháp giảm từ 48.678 xuống 15.034,25 giữa S0 và S1, tương ứng giảm 69,12%; số byte giảm từ 50.623.120 xuống 15.633.620. Đây là bằng chứng về tác động của synthetic gói điều khiển trực tiếp-only BHP trong cấu hình tái dựng, không phải tái lập nguyên trạng của mọi hành vi BHP giả mạo.', 412: 'Thứ hai, tái phân tích UCI404 so sánh bốn mô hình cơ sở có artifact đầy đủ trên 25 lượt đánh giá/mô hình. DecisionTree đạt macro-F1 0,8082 và MCC 0,6239; Flood Status là đặc trưng đơn mạnh nhất với độ chính xác 0,7106 và độ quan trọng hoán vị ngoài mẫu 0,2177. Dữ liệu có 860 hàng lặp và 215 vector dự báo duy nhất, cho thấy rủi ro phụ thuộc dữ liệu/biến đại diện; PSO–SVM và các kết quả gần hoàn hảo cũ chưa được tái lập.', 413: 'Thứ ba, chưa có bằng chứng đủ để lựa chọn mô hình vận hành theo đánh đổi Pareto giữa độ chính xác, độ trễ và tài nguyên. UCI404 có quy trình kiểm toán cho bốn mô hình cơ sở; PSO-SVM, độ trễ đầu cuối, kích thước mô hình và bộ phát hiện deployment artifact chưa được tái lập. Vì vậy, cây quyết định không được tuyên bố là bộ phát hiện vận hành đã kiểm chứng từ ma trận direct-BHP.', 414: 'Thứ tư, cơ chế kiểm soát BHP trực tiếp nguyên bản kiểm chứng được admission trước đặt trước tài nguyên, thứ tự nhân quả và hai cấu hình cơ cấu chấp hành. Tuy nhiên, đây là deterministic contemporaneous ngân sách token state machine; chưa phải closed-loop ML bộ phát hiện–decision–response với quy kết nguồn tin cậy, thống kê dương tính giả, danh sách xám, lùi thời gian/cơ chế trễ và cơ chế phản hồi và phục hồi đầy đủ.', 415: 'Thứ năm, ma trận 32 lượt bằng NS-2.35+nOBS trên bốn ô kịch bản và tám hạt giống ngẫu nhiên cố định cho thấy S2-rate-limit và S2-cách ly đều đạt 48.678 gói TCP hợp pháp, bằng 100% S0 và cao hơn S1. Kết quả chứng minh hiệu quả đáp ứng của mạng của hai cấu hình trong cấu hình nghiên cứu, không đồng nghĩa bộ phát hiện ML đã được nhúng trực tiếp vào NS-2 hoặc hệ thống đã sẵn sàng triển khai.', 418: 'Qua hai pha thực nghiệm và các gate kiểm toán, luận văn xác lập được những kết luận trong phạm vi artifact đã tái kiểm; những kết luận thiếu bằng chứng gốc không được xem là đã tái lập.', 419: 'Về bài toán phát hiện, UCI404 có mức lặp cao và chứa các biến có nguy cơ hậu nghiệm hoặc phản ánh cơ chế tạo nhãn. Giao thức nhóm bản sao làm kết quả giảm rõ rệt so với báo cáo cũ; do đó bộ chuẩn đánh giá này phù hợp cho kiểm toán dữ liệu ngoại tuyến, chưa phải bằng chứng bộ phát hiện vận hành.', 420: 'Về bài toán ứng phó, direct-BHP matrix cho thấy rate-limit và cách ly đều phục hồi byte TCP hợp pháp từ 15.633.620 ở S1 lên 50.623.120 ở S2, đạt 100% mức S0 trong cấu hình khảo sát. Đây là hiệu quả của oracle/cơ chế kiểm soát với đường xử lý điều khiển trực tiếp tổng hợp; chưa phải bằng chứng của bộ phát hiện-driven closed loop hoặc tính an toàn theo dương tính giả tổng quát.', 421: 'Về phương pháp đánh giá, luận văn tách riêng UCI404 kiểm toán ngoại tuyến, thực nghiệm mạng nguyên bản và bằng chứng đường xử lý BHP trực tiếp. Việc tách chuỗi nguồn gốc dữ liệu này ngăn không cho số liệu mô hình thử nghiệm tối thiểu tổng hợp hoặc benchmark cũ được dùng như bằng chứng NS-2.35+nOBS.', 424: 'Thứ nhất, bộ sinh BHP trực tiếp hiện là đường xử lý điều khiển không kèm chùm dữ liệu tổng hợp và còn dùng dấu mốc vòng đời nội bộ để chọn nhánh xử lý; cần chứng minh thêm rằng quyết định có thể dựa trên bằng chứng đồng thời quan sát được trên đường truyền, không bị ảnh hưởng bởi thông tin tiên nghiệm từ phía bộ sinh. Cần bổ sung end-to-end đặt trước tài nguyên outcome, tính nhất quán với nguồn sinh dữ liệu và tính an toàn của vòng đời xử lý.', 425: 'Thứ hai, cần kiểm thử khả năng đồng tồn tại giữa các gói điều khiển nOBS hợp pháp và gói điều khiển trực tiếps trên cùng điểm vào và định danh tin cậy, đồng thời đo ảnh hưởng ngoài mục tiêu theo luồng, độ trễ, độ dao động trễ và các lớp traffic khác. Zero collateral hiện tại chỉ giới hạn ở hai luồng TCP hợp pháp đã cấu hình.', 426: 'Thứ ba, cần mở rộng ma trận sang nhiều seed, rate, claimed chi phí đặt trước tài nguyên và tô-pô; đồng thời điều tra khả năng tái lập quá trình biên dịch, trình biên dịch và chuỗi công cụ và external nguồn gốc truy xuất cho run metadata. Tám seed hiện tại chỉ hỗ trợ mô tả trong tải công việc đã khai báo.', 427: 'Thứ tư, cần tích hợp bộ phát hiện trực tuyến nếu mục tiêu là closed-loop detection–decision–response: định nghĩa cửa sổ đặc trưng khả dụng theo quan hệ nhân quả, trusted quy kết nguồn, độ trễ p50/p95/p99 được đo, chuyển trạng thái, hoàn tác và kiểm thử tình huống lỗi. Không gọi guard hiện tại là PSO-SVM, RFC2698 đầy đủ hoặc bộ phát hiện ML vận hành.'}
for i,text in repl.items():
    setp(doc.paragraphs[i],text)

# Việt hóa có kiểm soát một số đoạn thân bài và Chương 3. Không thay thế
# chuỗi toàn cục: nhãn kịch bản, tên tệp, mã hành động và ký hiệu chuẩn phải
# được giữ nguyên để bảo đảm khả năng truy xuất và tái lập.
contextual_repl = {
    370: 'Công cụ mô phỏng là NS-2 phiên bản 2.35 kết hợp mô-đun nOBS trên môi trường Linux. Ma trận chính gồm 32 lượt chạy: bốn kịch bản nhân với tám hạt giống ngẫu nhiên cố định 101, 202, 303, 404, 505, 606, 707 và 808; mỗi lượt kéo dài 5 giây. Tô-pô tái dựng gồm bảy nút, liên kết quang 1000 Mb/s, hai luồng TCP hợp lệ với đường truy nhập 155 Mb/s và tám nguồn BHP trực tiếp. Mỗi nguồn phát BHP điều khiển không kèm chùm dữ liệu, với kích thước khai báo 1000 byte và một gói tin; tốc độ danh nghĩa là 12 Mb/s/nguồn, còn tốc độ hiệu dụng theo tám hạt giống nằm trong khoảng 9,8926–14,2687 Mb/s/nguồn, trung bình 12,3153 Mb/s/nguồn. Cơ chế kiểm soát dùng ngân sách số sự kiện và ngân sách chi phí đặt trước tài nguyên, không phải CIR payload 4 Mb/s. Các khoảng 95% trong phần này chỉ là khoảng mô tả trên tám hạt giống ngẫu nhiên cố định, không phải ước lượng bất định tổng quát.',
    381: 'Ma trận BHP trực tiếp không được dùng để tuyên bố đã xây dựng bộ chuẩn đánh giá học máy không suy biến. Các kết quả ML trên UCI và bộ chuẩn đánh giá nguyên bản theo cửa sổ được giữ ở nhánh kiểm toán riêng; trong đó PSO-SVM, độ trễ triển khai và bộ phát hiện trực tuyến chưa có tệp bằng chứng đầy đủ. Ma trận BHP trực tiếp chỉ kiểm chứng chuỗi xử lý nhân quả của tạo–quan sát–quyết định–tác động–kết quả của gói điều khiển và đáp ứng của mạng.',
    382: 'Các phép thử BHP trực tiếp kiểm tra rằng gói điều khiển được tạo, quan sát, quyết định và tác động theo đúng thứ tự thời gian; hành động DROP hoặc QUARANTINE xảy ra trước lần gọi thủ tục đặt trước tài nguyên tương ứng. Đây là tiêu chí kiểm định nhân quả của cơ chế kiểm soát, không phải phép đánh giá độ chính xác của một bộ phân loại.',
    384: 'Do đó, ma trận BHP trực tiếp được xem là thực nghiệm kiểm chứng kiểm soát chấp nhận và giảm thiểu ở mức mạng, không phải bộ chuẩn huấn luyện bộ phát hiện. Các kết luận chỉ giới hạn trong tô-pô bảy nút, cấu hình lưu lượng đã khai báo, thời lượng 5 giây và tám hạt giống ngẫu nhiên cố định.',
    387: 'Ma trận BHP trực tiếp không cung cấp đường cong khả năng phát hiện theo độ ẩn. Các số liệu MCC theo mức 1–35 Mb/s của bản cũ không có dấu vết thô và chuỗi nguồn gốc dữ liệu khớp với thí nghiệm mới, nên không được trình bày như kết quả tái lập.',
    389: 'Vì chưa có bộ chuẩn đánh giá bộ phát hiện trực tuyến không suy biến và chưa có phép đo độ trễ đầu cuối, luận văn không suy diễn từ ma trận này rằng một mô hình mạnh hơn là cần thiết ở vùng tấn công ẩn. Đây là giới hạn bằng chứng và là hướng thực nghiệm tiếp theo.',
    392: '3.6.5 Hiệu quả của hai chế độ kiểm soát BHP trực tiếp tại nút biên',
    393: 'Hai chế độ được hiện thực trong môi trường NS-2.35+nOBS nguyên bản tại nút biên. Cơ chế kiểm soát quan sát trực tiếp các gói điều khiển BHP, áp dụng ngân sách số sự kiện và ngân sách chi phí đặt trước tài nguyên trước lần đặt trước tài nguyên đầu tiên. S2-rate-limit cho phép một phần nhỏ gói điều khiển theo cấu hình rồi chuyển các gói điều khiển vượt cấu hình vào cách ly; S2-cách ly dùng cấu hình cách ly chặt hơn. Đây là đáp ứng của cơ chế kiểm soát dựa trên trạng thái và ngân sách thẻ bài với dữ liệu quan sát đồng thời, không phải việc nhúng trực tiếp cây quyết định hay mô hình PSO-SVM vào vòng chạy NS-2.',
    394: 'Trong ma trận BHP trực tiếp, S0 chỉ có hai luồng TCP hợp lệ; S1 bổ sung tám nguồn BHP trực tiếp không kèm chùm dữ liệu và dùng cơ chế kiểm soát đối chứng với ngân sách 1e9 nên không giới hạn thực tế; hai S2 giữ nguyên lưu lượng và hạt giống của S1 nhưng bật cấu hình kiểm soát tương ứng. BHP trực tiếp dùng dấu mốc vòng đời nội bộ để đi vào đường xử lý điều khiển tổng hợp; dấu mốc này không phải nhãn quan sát của bộ phát hiện. Do đó, kết quả đo hiệu quả giảm thiểu không chứng minh tỷ lệ dương tính giả đối với các gói điều khiển nOBS hợp lệ, vì các gói điều khiển hợp lệ không đi vào đường xử lý kiểm soát BHP trực tiếp tường minh.',
    395: 'Bảng 3.6 trình bày số gói TCP hợp lệ theo bốn ô kịch bản. Cả 32 lượt chạy đều kết thúc thành công; mỗi ô có tám hạt giống. Dấu vết thô out.tr được phân tích độc lập; chuỗi kiểm toán trong bhp_audit.log được bộ kiểm tra kiểm tra; kết quả được đối chiếu với validation.json và validation.rerun.json. Các khoảng trong bảng là khoảng mô tả trên tập hạt giống ngẫu nhiên cố định.',
    399: 'S2-rate-limit nâng số gói TCP hợp lệ từ 15.034,25 ở S1 lên 48.678, đạt 100% mức S0 và khôi phục toàn bộ phần TCP bị mất trong cấu hình khảo sát. S2-cách ly cũng đạt 48.678 gói, tương đương 100% mức S0 và khôi phục 100% phần bị mất. Hai chế độ có cùng kết quả vận hành mạng trong cả tám hạt giống; không có cơ sở kết luận chế độ nào ưu việt hơn về thông lượng.',
    401: 'Tổng hợp lại, ma trận 32 lượt chạy cho thấy trong cấu hình BHP trực tiếp đã khai báo, cả giới hạn tốc độ và cách ly đều ngăn phần lớn BHP điều khiển không kèm chùm dữ liệu trước đặt trước tài nguyên và phục hồi thông lượng TCP hợp lệ về đúng mức S0. Kết luận chỉ áp dụng cho tô-pô bảy nút, cấu hình lưu lượng, thời lượng 5 giây và tám hạt giống ngẫu nhiên cố định; không ghi nhận ảnh hưởng ngoài mục tiêu ở đây chỉ là kết quả của hai luồng được quan sát và phạm vi cơ chế kiểm soát đã khai báo, không phải tỷ lệ dương tính giả tổng quát.',
    404: 'Nhận định thứ nhất là sự phân biệt giữa các nhánh bằng chứng. UCI404 là kiểm toán ngoại tuyến của một bộ chuẩn đánh giá có nguy cơ trùng lặp, phụ thuộc dữ liệu và rò rỉ do cơ chế tạo nhãn hoặc biến đại diện; ma trận BHP trực tiếp nguyên bản là thực nghiệm chấp nhận và giảm thiểu ở mức mạng. Không được dùng kết quả của một nhánh để chứng minh kết luận của nhánh kia.',
    405: 'Nhận định thứ hai là bản chất đóng góp của luận văn. Đóng góp thực nghiệm hiện được hỗ trợ gồm: kiểm toán rò rỉ và tính tái lập của UCI404 theo giao thức có kiểm soát rò rỉ; xây dựng đường xử lý BHP trực tiếp nguyên bản có kiểm toán nhân quả; và đánh giá hai chế độ ứng phó tại nút biên bằng ma trận 32 ô. Ma trận mới cho thấy S1 giảm 69,12% byte TCP hợp lệ, còn cả hai S2 phục hồi 100% mức S0 trong cấu hình khảo sát.',
    406: 'Nhận định thứ ba là phạm vi và giới hạn suy luận. Bộ sinh BHP trực tiếp dùng dấu mốc vòng đời để chọn đường xử lý điều khiển tổng hợp; vì vậy kết quả không chứng minh bộ phát hiện suy luận từ bằng chứng quan sát được trên đường truyền, không phải tái lập nguyên trạng của PSO-SVM hay triển khai học máy vòng khép kín. Kiểm toán nguồn gốc truy xuất, tính nhất quán với nguồn sinh dữ liệu, khả năng đồng tồn tại của gói điều khiển hợp lệ và khả năng khái quát hóa trên nhiều tô-pô vẫn cần được mở rộng.',
    410: 'Ba kết quả được hỗ trợ trực tiếp là: tái phân tích UCI404 theo giao thức nhóm bản sao và chỉ ra rủi ro phụ thuộc dữ liệu/biến đại diện; xây dựng đường xử lý BHP trực tiếp không kèm chùm dữ liệu trong môi trường NS-2.35+nOBS nguyên bản với kiểm toán thứ tự trước đặt trước tài nguyên; và đánh giá hai cấu hình kiểm soát bằng ma trận 32 lượt chạy. Bộ chuẩn đánh giá học máy nguyên bản không suy biến, lựa chọn mô hình vận hành và vòng khép kín do bộ phát hiện điều khiển chưa đạt tiêu chí kiểm định bằng chứng.',
    412: 'Thứ hai, tái phân tích UCI404 so sánh bốn mô hình cơ sở có tệp bằng chứng đầy đủ trên 25 lượt đánh giá/mô hình. DecisionTree đạt macro-F1 0,8082 và MCC 0,6239; Flood Status là đặc trưng đơn mạnh nhất với độ chính xác 0,7106 và độ quan trọng hoán vị ngoài mẫu 0,2177. Dữ liệu có 860 hàng lặp và 215 véc-tơ dự báo duy nhất, cho thấy rủi ro phụ thuộc dữ liệu/biến đại diện; PSO–SVM và các kết quả gần hoàn hảo cũ chưa được tái lập.',
    413: 'Thứ ba, chưa có bằng chứng đủ để lựa chọn mô hình vận hành theo đánh đổi Pareto giữa độ chính xác, độ trễ và tài nguyên. UCI404 có quy trình kiểm toán cho bốn mô hình cơ sở; PSO-SVM, độ trễ đầu cuối, kích thước mô hình và tệp bằng chứng triển khai của bộ phát hiện chưa được tái lập. Vì vậy, cây quyết định không được tuyên bố là bộ phát hiện vận hành đã kiểm chứng từ ma trận BHP trực tiếp.',
    414: 'Thứ tư, cơ chế kiểm soát BHP trực tiếp nguyên bản kiểm chứng được chấp nhận trước đặt trước tài nguyên, thứ tự nhân quả và hai cấu hình cơ cấu chấp hành. Tuy nhiên, đây là máy trạng thái xác định dựa trên quan sát đồng thời và ngân sách thẻ bài; chưa phải vòng khép kín gồm bộ phát hiện học máy–quyết định–ứng phó với quy kết nguồn tin cậy, thống kê dương tính giả, danh sách xám, lùi thời gian/cơ chế trễ và cơ chế phản hồi, phục hồi đầy đủ.',
    418: 'Qua hai pha thực nghiệm và các tiêu chí kiểm định, luận văn xác lập được những kết luận trong phạm vi tệp bằng chứng đã tái kiểm; những kết luận thiếu bằng chứng gốc không được xem là đã tái lập.',
    420: 'Về bài toán ứng phó, ma trận BHP trực tiếp cho thấy giới hạn tốc độ và cách ly đều phục hồi byte TCP hợp lệ từ 15.633.620 ở S1 lên 50.623.120 ở S2, đạt 100% mức S0 trong cấu hình khảo sát. Đây là hiệu quả của cơ chế kiểm soát có thông tin đầu vào được xác định trước trong thí nghiệm với đường xử lý điều khiển trực tiếp tổng hợp; chưa phải bằng chứng của vòng khép kín do bộ phát hiện điều khiển hoặc tính an toàn theo dương tính giả tổng quát.',
    421: 'Về phương pháp đánh giá, luận văn tách riêng kiểm toán ngoại tuyến UCI404, thực nghiệm mạng nguyên bản và bằng chứng đường xử lý BHP trực tiếp. Việc tách chuỗi nguồn gốc dữ liệu này ngăn không cho số liệu mô hình thử nghiệm tối thiểu tổng hợp hoặc bộ chuẩn đánh giá cũ được dùng như bằng chứng NS-2.35+nOBS.',
    425: 'Thứ hai, cần kiểm thử khả năng đồng tồn tại giữa các gói điều khiển nOBS hợp lệ và gói điều khiển trực tiếp trên cùng điểm vào và định danh tin cậy, đồng thời đo ảnh hưởng ngoài mục tiêu theo luồng, độ trễ, độ dao động trễ và các lớp lưu lượng khác. Hiện chưa ghi nhận ảnh hưởng ngoài mục tiêu chỉ giới hạn ở hai luồng TCP hợp lệ đã cấu hình.',
    426: 'Thứ ba, cần mở rộng ma trận sang nhiều hạt giống, tốc độ, chi phí đặt trước tài nguyên đã khai báo và tô-pô; đồng thời điều tra khả năng tái lập quá trình biên dịch, trình biên dịch, chuỗi công cụ và nguồn gốc truy xuất bên ngoài cho siêu dữ liệu lượt chạy. Tám hạt giống hiện tại chỉ hỗ trợ mô tả trong tải công việc đã khai báo.',
    427: 'Thứ tư, cần tích hợp bộ phát hiện trực tuyến nếu mục tiêu là vòng khép kín phát hiện–quyết định–ứng phó: định nghĩa cửa sổ đặc trưng khả dụng theo quan hệ nhân quả, quy kết nguồn tin cậy, độ trễ p50/p95/p99 được đo, chuyển trạng thái, hoàn tác và kiểm thử tình huống lỗi. Không gọi cơ chế kiểm soát hiện tại là PSO-SVM, RFC 2698 đầy đủ hoặc bộ phát hiện học máy vận hành.',
}
for i, text in contextual_repl.items():
    setp(doc.paragraphs[i], text)
# Update body captions only. List-of-tables/list-of-figures entries contain a
# tab + rendered page number and are set explicitly in `repl`; touching them
# here would silently strip their page numbers.
for i,p in enumerate(doc.paragraphs):
    if '\t' in p.text:
        continue
    if 'Bảng 3.2. Ba kịch bản' in p.text: setp(p,p.text.replace('Ba kịch bản','Bốn kịch bản'))
    if 'Hình 3.4.' in p.text: setp(p,'Hình 3.4. Tác động của BHP điều khiển trực tiếp không kèm chùm dữ liệu lên thông lượng TCP hợp pháp giữa S0 và S1 trên tám hạt giống ngẫu nhiên cố định.')
    if 'Hình 3.7.' in p.text: setp(p,'Hình 3.7. Hiệu quả của hai cấu hình kiểm soát BHP trực tiếp trên tám hạt giống ngẫu nhiên cố định.')
    if 'Bảng 3.6.' in p.text: setp(p,'Bảng 3.6. Thông lượng TCP hợp pháp theo bốn kịch bản direct-BHP (khoảng mô tả trên tám seed).')
# Robust cleanup for legacy paragraphs whose positional indices can shift when
# tables, captions, or embedded drawings are read by python-docx.  Matching by
# stable paragraph prefixes prevents stale claims from surviving a rebuild.
legacy_prefix_replacements = {
    'Bảng 3.1. Kết quả năm mô hình học máy': 'Bảng 3.1. Kết quả tái phân tích bốn mô hình cơ sở trên UCI404 bằng StratifiedGroupKFold (25 lượt đánh giá/mô hình).',
    'Bảng 3.2. Ba kịch bản mô phỏng đối chứng.': 'Bảng 3.2. Bốn kịch bản mô phỏng đối chứng.',
    'Hình 3.3. Kiến trúc cơ chế phát hiện--ứng phó khép kín': 'Hình 3.3. Kiến trúc mục tiêu phát hiện–quyết định–ứng phó tại nút biên; ma trận hiện tại chỉ kiểm chứng cơ chế kiểm soát BHP trực tiếp trước khi đặt trước tài nguyên và hai cấu hình ứng phó.',
    'Hình 3.7. Hiệu quả của cơ chế ứng phó khép kín': 'Hình 3.7. Hiệu quả của hai cấu hình kiểm soát BHP trực tiếp trên tám hạt giống ngẫu nhiên cố định.',
    '(5) Mô phỏng kịch bản tấn công bằng NS2/OMNeT++': '(5) Đánh giá tác động của BHP điều khiển trực tiếp và hai cấu hình kiểm soát bằng môi trường NS-2.35 kết hợp mô-đun nOBS nguyên bản; các kết quả được giới hạn trong tô-pô, lưu lượng, thời lượng và hạt giống đã khai báo.',
    '**Mục tiêu tổng quát: Nghiên cứu và đề xuất một cơ chế tích hợp khép kín': '**Mục tiêu tổng quát: Làm rõ nguy cơ của tấn công ngập lụt BHP và đánh giá cơ chế kiểm soát, ứng phó tại nút biên trong phạm vi bằng chứng có thể tái kiểm.',
    '(4) Đề xuất và tích hợp cơ chế ứng phó': '(4) Đề xuất kiến trúc phát hiện–quyết định–ứng phó và hiện thực, đánh giá hai cấu hình kiểm soát xác định tại nút biên; phần phát hiện học máy trực tuyến và vòng phản hồi đầy đủ được xác định là hướng mở rộng.',
    'Ở pha thứ nhất, luận văn huấn luyện và so sánh năm mô hình học máy': 'Ở pha thứ nhất, luận văn tái phân tích bốn mô hình cơ sở có mã và đầu ra đầy đủ gồm DecisionTree, SVM-RBF, KNN và Gaussian Naïve Bayes. PSO–SVM chỉ được giữ ở phần tổng quan đối chiếu vì chưa có artifact gốc đủ để tái lập.',
    'Bảng 3.1 trình bày kết quả chạy thật của năm mô hình': 'Bảng 3.1 trình bày kết quả trung bình của 25 lượt đánh giá theo giao thức có kiểm soát nhóm bản sao. DecisionTree đạt macro-F1 0,8082 và MCC 0,6239; SVM-RBF đạt macro-F1 0,7680 và MCC 0,6095. Không mô hình nào đạt kết quả hoàn hảo.',
    '### **3.3.2 Kết quả năm mô hình': '### **3.3.2 Kết quả tái phân tích bốn mô hình',
    'Pha thứ nhất tập trung so sánh các mô hình học máy trên bộ dữ liệu chuẩn': 'Pha thứ nhất tập trung tái phân tích bốn mô hình cơ sở trên bộ dữ liệu chuẩn, đồng thời kiểm định tính toàn vẹn của dữ liệu trước khi diễn giải điểm số. Quy trình này trả lời câu hỏi liệu kết quả có phụ thuộc vào bản sao hoặc đặc trưng liên quan đến cơ chế tạo nhãn hay không.',
    'Sau khi đã chỉ ra rằng bộ dữ liệu UCI bị suy biến': 'Sau khi kiểm tra rủi ro phụ thuộc dữ liệu của UCI404, luận văn chuyển sang thực nghiệm native NS-2.35+nOBS để đánh giá tác động mạng và cơ chế kiểm soát BHP trực tiếp. Hai pha không dùng kết quả của pha này để thay thế bằng chứng của pha kia.',
    '### **3.5 Thiết kế cơ chế ứng phó khép kín': '### **3.5 Kiến trúc mục tiêu và cơ chế kiểm soát đã hiện thực',
    'Trên cơ sở đó, luận văn thiết kế một vòng ứng phó khép kín': 'Trên cơ sở đó, luận văn mô tả kiến trúc mục tiêu gồm dữ liệu giám sát, phát hiện, quyết định, cơ cấu chấp hành và phản hồi. Artifact hiện tại chỉ hiện thực cơ chế kiểm soát xác định trước khi đặt trước tài nguyên và hai cấu hình ứng phó tại nút biên.',
    '### **3.5.2 Mô hình mạng và ba kịch bản': '### **3.5.2 Mô hình mạng và bốn kịch bản đối chứng',
    'Để đánh giá có đối chứng, luận văn thiết lập ba kịch bản.': 'Để đánh giá có đối chứng, ma trận thiết lập bốn kịch bản: S0 nền không tấn công; S1 có BHP trực tiếp với guard đối chứng không giới hạn thực tế; S2-rate-limit bật giới hạn theo ngân sách; và S2-isolation bật cách ly tại nút biên.',
    'Cấu trúc ba kịch bản này được tóm tắt trong': 'Cấu trúc bốn kịch bản này được tóm tắt trong',
    'Đóng góp khoa học cốt lõi của luận văn không nằm ở việc đề xuất một thuật toán phát hiện mới': 'Đóng góp phương pháp luận của luận văn nằm ở việc tách và kiểm toán hai nhánh bằng chứng: (i) kiểm tra độ tin cậy của đánh giá UCI404 dưới giao thức nhóm bản sao; và (ii) xây dựng ma trận native direct-BHP có kiểm toán nhân quả để đo tác động mạng và hiệu quả của hai cấu hình kiểm soát. Luận văn không tuyên bố đã tái lập PSO-SVM hoặc hoàn thiện hệ thống học máy vòng khép kín.',
    'Bộ dữ liệu được sử dụng ở pha thứ nhất là tập \\"OBS-Network-DataSet\\"': 'Bộ dữ liệu được sử dụng ở pha thứ nhất là tập "OBS-Network-DataSet" trên UCI Machine Learning Repository, truy xuất từ nguồn được ghi trong artifact cấu hình là https://example.com. Artifact ARFF được lưu tại data/uci404/extracted/OBS-Network-DataSet_2_Aug27.arff và được kiểm soát bằng SHA-256; dữ liệu gồm 1.075 mẫu, 21 biến dự báo và nhãn Class bốn lớp.',
}
for prefix, replacement in legacy_prefix_replacements.items():
    for p in doc.paragraphs:
        if p.text.startswith(prefix):
            setp(p, replacement)

def replace_table_rows(table, rows):
    """Resize a table and replace every cell from a validated row matrix."""
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    target_cols = len(rows[0])
    assert all(len(row) == target_cols for row in rows), 'inconsistent replacement-table width'
    for ri, row in enumerate(rows):
        tr = table.rows[ri]._tr
        while len(table.rows[ri].cells) < target_cols:
            tr.add_tc()
        while len(tr.tc_lst) > target_cols:
            tr.remove(tr.tc_lst[-1])
        current_cells = table.rows[ri].cells
        assert len(current_cells) == target_cols, (
            f'row {ri} has {len(current_cells)} cells after resize; expected {target_cols}'
        )
        for ci, val in enumerate(row):
            current_cells[ci].text = val


def repeat_table_header(table):
    """Repeat the first row when Word/LibreOffice splits a table across pages."""
    trPr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = trPr.find(qn('w:tblHeader'))
    if tbl_header is None:
        tbl_header = OxmlElement('w:tblHeader')
        trPr.append(tbl_header)
    tbl_header.set(qn('w:val'), 'true')

rows2=[
    ['Mô hình','Accuracy','Balanced accuracy','Macro-F1','MCC','Số lượt đánh giá','Độ trễ suy luận'],
    ['DecisionTree','0,7523','0,8171','0,8082','0,6239','25','Chưa đo'],
    ['SVM-RBF','0,7404','0,7748','0,7680','0,6095','25','Chưa đo'],
    ['KNN','0,6492','0,6734','0,6573','0,4750','25','Chưa đo'],
    ['GaussianNB','0,6742','0,7581','0,6838','0,5526','25','Chưa đo'],
]
replace_table_rows(doc.tables[2], rows2)
# Keep Table 3.1 together and make the model-name column readable. The source
# layout left too little room at the bottom of the page, splitting the table
# after its first data row and wrapping "DecisionTree" mid-word.
doc.paragraphs[345].paragraph_format.page_break_before = True
doc.paragraphs[345].paragraph_format.keep_with_next = True
table31 = doc.tables[2]
table31.autofit = False
table31_widths = (1.25, 0.72, 1.05, 0.72, 0.62, 0.85, 0.85)
from docx.shared import Inches, Pt
for row in table31.rows:
    trPr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement('w:cantSplit')
    trPr.append(cant_split)
    for ci, width in enumerate(table31_widths):
        row.cells[ci].width = Inches(width)
        for paragraph in row.cells[ci].paragraphs:
            paragraph.paragraph_format.keep_together = True
            for run in paragraph.runs:
                run.font.size = Pt(9)
rows3=[['Kịch bản','Lưu lượng hợp pháp','BHP trực tiếp','Biện pháp ứng phó'],['S0 – nền','Có','Không','Không'],['S1 – tấn công','Có','Có, cấu hình đối chứng không giới hạn','Không'],['S2-rate-limit','Có','Có','Ngân sách sự kiện/đặt trước tài nguyên'],['S2-cách ly','Có','Có','Cách ly']]
replace_table_rows(doc.tables[3], rows3)
rows4=[
    ['Chỉ số','S0 – nền','S1 – tấn công','Thay đổi S1/S0','Phạm vi'],
    ['Gói TCP hợp pháp','48.678','15.034,25','−69,12%','8 hạt giống ngẫu nhiên cố định'],
    ['Byte TCP hợp pháp','50.623.120','15.633.620','−69,12%','8 hạt giống ngẫu nhiên cố định'],
    ['Cặp burst quang','4.917','2.060,5','−58,09%','8 hạt giống ngẫu nhiên cố định'],
    ['Reservation thành công','19.606','8.216,5','−58,09%','8 hạt giống ngẫu nhiên cố định'],
    ['Số lượt loại bỏ dữ liệu tường minh','0','0','—','8 hạt giống ngẫu nhiên cố định'],
]
replace_table_rows(doc.tables[4], rows4)
rows5=[
    ['Hạng mục kiểm toán','Artifact hiện có','Kết quả gate','Kết luận được phép'],
    ['Nguồn native','320 cửa sổ từ 16 cell S0/S1','Có trace và nhóm theo seed','Chỉ dùng để kiểm toán pipeline'],
    ['Không suy biến','7 đặc trưng đơn đạt 1,0','Không đạt','Không báo cáo năng lực bộ phát hiện'],
    ['Tổng quát hóa','Một topology và nhãn đồng nhất scenario','Không đạt','Không suy rộng sang vận hành'],
    ['Detector online','Không có prediction/action linkage','Không đạt','Ma trận BHP trực tiếp không phải bộ chuẩn đánh giá học máy'],
]
replace_table_rows(doc.tables[5], rows5)
rows6=[
    ['Thành phần cần có','Trạng thái artifact','Kết luận'],
    ['Native traces theo nhiều mức attack rate','Chưa có chuỗi nguồn gốc dữ liệu khớp','Chưa lập đường cong'],
    ['Grouped holdout theo run/rate/topology','Chưa có','Chưa đánh giá tổng quát hóa'],
    ['Raw fold predictions và uncertainty','Chưa có','Không tái lập MCC cũ'],
    ['Latency feature→decision→action','Chưa đo','Không kết luận thời gian thực'],
]
replace_table_rows(doc.tables[6], rows6)
rows7=[['Kịch bản','n','Gói TCP hợp pháp (TB)','Khoảng mô tả','So với S0'],['S0 – nền','8','48.678','[48.678; 48.678]','100%'],['S1 – tấn công','8','15.034,25','[13.134,818; 16.933,682]','30,88%'],['S2-rate-limit','8','48.678','[48.678; 48.678]','100%'],['S2-cách ly','8','48.678','[48.678; 48.678]','100%']]
replace_table_rows(doc.tables[7], rows7)
for table in doc.tables:
    repeat_table_header(table)

# Normalize quotation marks in the bibliography to one IEEE-style form.
# Only reference paragraphs are touched; apostrophes and other punctuation stay intact.
for p in doc.paragraphs:
    if not p.text.lstrip().startswith('['):
        continue
    text = p.text
    if '"' not in text:
        continue
    parts = text.split('"')
    normalized = ''.join(
        ('“' if i % 2 else '”') + part if i else part
        for i, part in enumerate(parts)
    )
    setp(p, normalized)
# Load canonical native evidence once. All numerical figures and tables below
# must derive from these files; hard-coded plotting values are forbidden.
canonical_summary_path = base/'evidence/direct_bhp_matrix/summary.json'
canonical_per_seed_path = base/'evidence/direct_bhp_matrix/per_seed.csv'
with canonical_summary_path.open(encoding='utf-8') as stream:
    canonical_summary = json.load(stream)
with canonical_per_seed_path.open(encoding='utf-8', newline='') as stream:
    canonical_rows = list(csv.DictReader(stream))
assert len(canonical_rows) == 32, f'expected 32 native rows, got {len(canonical_rows)}'


def mean(label, metric):
    return canonical_summary['summary'][label][metric]['mean']


s0_packets = mean('S0', 'legal_tcp_packets')
s1_packets = mean('S1', 'legal_tcp_packets')
s2_rate_packets = mean('S2_rate_limit', 'legal_tcp_packets')
s2_isolation_packets = mean('S2_isolation', 'legal_tcp_packets')

# Replace embedded legacy figures with direct-matrix figures if relationship members exist.
fig1=base/'docx_work/figure_3_4_direct.png'; fig2=base/'docx_work/figure_3_7_direct.png'
fig5=base/'docx_work/figure_3_5_scope.png'; fig6=base/'docx_work/figure_3_6_evidence_gap.png'
# generate simple figures from validated means; avoid dependency on old scripts.
import matplotlib.pyplot as plt
plt.rcParams.update({'font.family':'DejaVu Sans','font.size':10})
# Hình 3.4 và 3.7: số liệu đọc trực tiếp từ summary.json canonical đã kiểm định.
fig,ax=plt.subplots(figsize=(7.1,5.1),dpi=180); bars=ax.bar(['S0 – nền','S1 – BHP trực tiếp'],[s0_packets,s1_packets],color=['#2F6B9A','#C44E52'],width=.58); ax.set_ylabel('Số gói TCP hợp pháp nhận được'); ax.set_title('Tác động của BHP điều khiển trực tiếp'); ax.grid(axis='y',alpha=.25); ax.bar_label(bars,labels=[f'{s0_packets:,.0f}'.replace(',','.'),f'{s1_packets:,.1f}'.replace(',','X').replace('.',',').replace('X','.')],padding=5,fontweight='bold'); ax.text(.5,-.15,'Trung bình trên 8 hạt giống ngẫu nhiên cố định',transform=ax.transAxes,ha='center',fontsize=9); fig.tight_layout(); fig.savefig(fig1,bbox_inches='tight'); plt.close(fig)
fig,ax=plt.subplots(figsize=(7.1,5.1),dpi=180); plot_means=[s0_packets,s1_packets,s2_rate_packets,s2_isolation_packets]; bars=ax.bar(['S0','S1','S2-giới hạn tốc độ','S2-cách ly'],plot_means,color=['#2F6B9A','#C44E52','#55A868','#8172B3'],width=.65); ax.set_ylabel('Số gói TCP hợp pháp nhận được'); ax.set_title('So sánh bốn kịch bản BHP trực tiếp'); ax.grid(axis='y',alpha=.25); ax.bar_label(bars,labels=[f'{v:,.1f}'.replace(',','X').replace('.',',').replace('X','.').removesuffix(',0') for v in plot_means],padding=5,fontsize=9,fontweight='bold'); ax.text(.5,-.15,'Trung bình trên 8 hạt giống; hai S2 bằng mức S0',transform=ax.transAxes,ha='center',fontsize=8.5); fig.tight_layout(); fig.savefig(fig2,bbox_inches='tight'); plt.close(fig)
# Hình 3.5 và 3.6 dùng trực tiếp số liệu từng lượt chạy từ per_seed.csv,
# thay cho sơ đồ chữ. Bất kỳ lỗi thiếu cột hay thiếu lượt chạy đều làm dừng bản dựng.
rows=canonical_rows
seeds=sorted({int(row['seed']) for row in rows})
labels=['S0','S1','S2_rate_limit','S2_isolation']
label_vi={'S0':'S0 – nền','S1':'S1 – BHP trực tiếp','S2_rate_limit':'S2 – giới hạn tốc độ','S2_isolation':'S2 – cách ly'}
colors={'S0':'#2F6B9A','S1':'#C44E52','S2_rate_limit':'#55A868','S2_isolation':'#8172B3'}
by={(int(row['seed']),row['label']):row for row in rows}
assert all((seed,label) in by for seed in seeds for label in labels)
fig,ax=plt.subplots(figsize=(7.1,4.9),dpi=180)
for label in labels:
    values=[int(by[(seed,label)]['legal_tcp_packets']) for seed in seeds]
    ax.plot(seeds,values,marker='o',linewidth=2,markersize=4,label=label_vi[label],color=colors[label])
ax.set_xlabel('Hạt giống ngẫu nhiên cố định'); ax.set_ylabel('Số gói TCP hợp pháp nhận được')
ax.set_title('Kết quả theo từng lượt chạy native'); ax.grid(alpha=.25); ax.set_xticks(seeds); ax.legend(loc='lower left',fontsize=8); fig.tight_layout(); fig.savefig(fig5,bbox_inches='tight'); plt.close(fig)
fig,ax=plt.subplots(figsize=(7.1,4.9),dpi=180)
action_labels=['Được chấp nhận','Giải phóng có kiểm soát','Loại do vượt cấu hình','Cách ly']
actions=['allow','release','drop_over_profile','quarantine']
profiles=['S1','S2_rate_limit','S2_isolation']
profile_vi=['S1 – đối chứng','S2 – giới hạn tốc độ','S2 – cách ly']
bottom=[0]*len(profiles)
for action,action_vi,color in zip(actions,action_labels,['#4C78A8','#72B7B2','#F2CF5B','#E45756']):
    values=[sum(int(by[(seed,label)][action]) for seed in seeds) for label in profiles]
    ax.bar(profile_vi,values,bottom=bottom,label=action_vi,color=color,width=.62)
    bottom=[a+b for a,b in zip(bottom,values)]
ax.set_ylabel('Số quyết định BHP, tổng trên 8 hạt giống'); ax.set_title('Dấu vết quyết định của các cấu hình ứng phó'); ax.legend(fontsize=8,loc='upper right'); ax.grid(axis='y',alpha=.22); fig.tight_layout(); fig.savefig(fig6,bbox_inches='tight'); plt.close(fig)

def replace_media(path,mapping):
    tmp=path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(path) as zin, zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist(): zout.writestr(item, mapping.get(item.filename, zin.read(item.filename)))
    os.replace(tmp,path)
doc.save(out)
with zipfile.ZipFile(out) as z:
    media = [n for n in z.namelist() if n.startswith("word/media/")]
media_updates={}
uci_figs=base/'source_only/uci404/outputs/figures'
for name, source in {
    'word/media/image24.png': uci_figs/'single_feature_audit.png',
    'word/media/image25.png': uci_figs/'rf_oof_permutation_importance.png',
    'word/media/image26.png': fig1,
    'word/media/image27.png': fig5,
    'word/media/image28.png': fig6,
    'word/media/image29.png': fig2,
}.items():
    if name in media and source.is_file():
        media_updates[name]=source.read_bytes()
replace_media(out, media_updates)

# Emit a figure-lineage manifest for the exact media embedded in the rebuilt
# DOCX. Hình 3.3 is explicitly architectural; every other Chapter 3 figure is
# tied to measured pipeline/native outputs.
def sha256_path(path):
    return hashlib.sha256(path.read_bytes()).hexdigest()

with zipfile.ZipFile(out) as z:
    embedded_hashes = {
        name: hashlib.sha256(z.read(name)).hexdigest()
        for name in z.namelist() if name.startswith('word/media/')
    }
lineage = {
    'schema': 'thesis-figure-lineage-v1',
    'document': {'path': str(out.relative_to(base)), 'sha256': sha256_path(out)},
    'generator': {'path': str(Path(__file__).resolve().relative_to(base)), 'sha256': sha256_path(Path(__file__).resolve())},
    'figures': [
        {
            'figure': 'Hình 3.1', 'kind': 'experimental-uci404',
            'output': str((uci_figs/'single_feature_audit.png').relative_to(base.parent)),
            'output_sha256': sha256_path(uci_figs/'single_feature_audit.png'),
            'embedded_member': 'word/media/image24.png',
            'embedded_sha256': embedded_hashes.get('word/media/image24.png'),
            'inputs': [
                {'path': str((uci_figs.parent/'raw/single_feature_fold_metrics.csv').relative_to(base.parent)), 'sha256': sha256_path(uci_figs.parent/'raw/single_feature_fold_metrics.csv')},
                {'path': str((uci_figs.parent/'summary/single_feature_summary.csv').relative_to(base.parent)), 'sha256': sha256_path(uci_figs.parent/'summary/single_feature_summary.csv')},
            ],
        },
        {
            'figure': 'Hình 3.2', 'kind': 'experimental-uci404',
            'output': str((uci_figs/'rf_oof_permutation_importance.png').relative_to(base.parent)),
            'output_sha256': sha256_path(uci_figs/'rf_oof_permutation_importance.png'),
            'embedded_member': 'word/media/image25.png',
            'embedded_sha256': embedded_hashes.get('word/media/image25.png'),
            'inputs': [
                {'path': str((uci_figs.parent/'raw/rf_oof_permutation_fold.csv').relative_to(base.parent)), 'sha256': sha256_path(uci_figs.parent/'raw/rf_oof_permutation_fold.csv')},
                {'path': str((uci_figs.parent/'summary/rf_permutation_importance_summary.csv').relative_to(base.parent)), 'sha256': sha256_path(uci_figs.parent/'summary/rf_permutation_importance_summary.csv')},
            ],
        },
        {
            'figure': 'Hình 3.3', 'kind': 'conceptual-architecture',
            'output': None, 'output_sha256': None,
            'embedded_member': 'word/media/image23.png',
            'embedded_sha256': embedded_hashes.get('word/media/image23.png'),
            'inputs': [],
            'note': 'Sơ đồ kiến trúc khái niệm; không được trình bày như kết quả thực nghiệm.',
        },
        *[
            {
                'figure': figure, 'kind': 'experimental-native-direct-bhp',
                'output': str(path.relative_to(base)), 'output_sha256': sha256_path(path),
                'embedded_member': member, 'embedded_sha256': embedded_hashes.get(member),
                'inputs': [
                    {'path': str(canonical_summary_path.relative_to(base)), 'sha256': sha256_path(canonical_summary_path)},
                    {'path': str(canonical_per_seed_path.relative_to(base)), 'sha256': sha256_path(canonical_per_seed_path)},
                ],
            }
            for figure, path, member in [
                ('Hình 3.4', fig1, 'word/media/image26.png'),
                ('Hình 3.5', fig5, 'word/media/image27.png'),
                ('Hình 3.6', fig6, 'word/media/image28.png'),
                ('Hình 3.7', fig2, 'word/media/image29.png'),
            ]
        ],
    ],
}
for entry in lineage['figures']:
    if entry['output_sha256'] is not None:
        assert entry['embedded_sha256'] == entry['output_sha256'], entry['figure']
lineage_path = base/'docx_work/FIGURE_LINEAGE.json'
lineage_path.write_text(json.dumps(lineage, ensure_ascii=False, indent=2) + '\n', encoding='utf-8')
# Create the review copy against the immutable original DOCX, not against the
# already-updated packaged source. Every changed paragraph/cell is marked. A
# changed embedded image is marked by shading its containing drawing paragraph.
new=Document(out); old=Document(original)
oldp=[p.text for p in old.paragraphs]; oldt=[[[c.text for c in r.cells] for r in t.rows] for t in old.tables]

def shade_paragraph(p):
    for r in p.runs:
        r.font.highlight_color=WD_COLOR_INDEX.YELLOW
    pPr=p._p.get_or_add_pPr()
    shd=pPr.find(qn('w:shd'))
    if shd is None:
        shd=OxmlElement('w:shd'); pPr.append(shd)
    shd.set(qn('w:fill'),'FFF2CC')

for i,p in enumerate(new.paragraphs):
    ov=oldp[i] if i<len(oldp) else None
    if p.text != ov:
        shade_paragraph(p)
for ti,t in enumerate(new.tables):
    old_table = oldt[ti] if ti < len(oldt) else []
    new_table = [[c.text for c in r.cells] for r in t.rows]
    if new_table != old_table:
        # Mark the complete table when its shape or any cell changes. This also
        # covers newly appended rows, removed/reordered rows, and merged-cell
        # layouts that cannot be audited reliably by positional cell matching.
        for row in t.rows:
            for c in row.cells:
                for p in c.paragraphs:
                    shade_paragraph(p)
                tcPr=c._tc.get_or_add_tcPr()
                shd=tcPr.find(qn('w:shd'))
                if shd is None:
                    shd=OxmlElement('w:shd'); tcPr.append(shd)
                shd.set(qn('w:fill'),'FFF2CC')

def media_hashes(path):
    with zipfile.ZipFile(path) as z:
        return {
            n: hashlib.sha256(z.read(n)).hexdigest()
            for n in z.namelist()
            if n.startswith("word/media/")
        }

old_media, new_media = media_hashes(original), media_hashes(out)
changed_media={n for n in set(old_media)|set(new_media) if old_media.get(n)!=new_media.get(n)}
for p in new.paragraphs:
    for blip in p._p.xpath('.//a:blip'):
        rid=blip.get(qn('r:embed'))
        if rid and rid in new.part.rels:
            target='word/'+str(new.part.rels[rid].target_ref).replace('\\','/').lstrip('/')
            if target in changed_media:
                shade_paragraph(p)
                break
new.save(high)
print(out); print(high)

# NOTE: DOCX field/TOC/citation polishing (fix_toc_lists.py, fix_citations.py)
# is a personal manuscript-formatting pass kept local only; it is intentionally
# not tracked in this repository and not auto-invoked from this build step,
# to keep the published source focused on the direct-BHP experiment.

