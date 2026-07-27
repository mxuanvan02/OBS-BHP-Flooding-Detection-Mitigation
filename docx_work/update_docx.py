from pathlib import Path
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import zipfile, shutil, tempfile, os

BASE=Path(__file__).resolve().parents[1]
src=BASE/'deliverables/LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
out=BASE/'LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
work=BASE/'docx_work'

def setp(p, text):
    # Preserve paragraph-level style/alignment; keep first run's formatting where possible.
    if p.runs:
        first=p.runs[0]
        first.text=text
        for r in p.runs[1:]: r.text=''
    else:
        p.add_run(text)

def replace_exact(old, new):
    hits=0
    for p in doc.paragraphs:
        if p.text.strip()==old:
            setp(p,new); hits+=1
    return hits

def replace_starts(prefix, new):
    hits=0
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            setp(p,new); hits+=1
    return hits

# Generate replacement figures.
plt.rcParams.update({'font.family':'DejaVu Sans','font.size':10})
# Figure 3.4: S0/S1 with fixed-seed descriptive intervals
fig,ax=plt.subplots(figsize=(7.1,5.1),dpi=180)
labels=['S0 – nền','S1 – tấn công']
means=[3426.0,316.25]
lo=[3426.0,278.1379500645902]; hi=[3426.0,354.3620499354098]
yerr=np.array([[m-l for m,l in zip(means,lo)],[h-m for m,h in zip(means,hi)]])
bars=ax.bar(labels,means,color=['#2F6B9A','#C44E52'],width=.58,yerr=yerr,capsize=6)
ax.set_ylabel('Gói TCP hợp pháp nhận được')
ax.set_title('Tác động của tấn công ở tải danh nghĩa 40 Mb/s/nguồn')
ax.grid(axis='y',alpha=.25)
ax.bar_label(bars,labels=['3.426','316,25'],padding=5,fontweight='bold')
ax.text(.5,-.16,'Thanh sai số: khoảng 95% mô tả trên 8 seed cố định',transform=ax.transAxes,ha='center',fontsize=9)
fig.tight_layout(); fig.savefig(work/'figure_3_4_new.png',bbox_inches='tight'); plt.close(fig)

fig,ax=plt.subplots(figsize=(7.1,5.1),dpi=180)
labels=['S0\nNền','S1\nTấn công','S2\nGiới hạn tốc độ','S2\nCách ly']
means=[3426.0,316.25,2823.875,2855.625]
lo=[3426.0,278.1379500645902,2814.5073278827035,2686.258787979667]
hi=[3426.0,354.3620499354098,2833.2426721172965,3024.991212020333]
yerr=np.array([[m-l for m,l in zip(means,lo)],[h-m for m,h in zip(means,hi)]])
colors=['#2F6B9A','#C44E52','#55A868','#8172B3']
bars=ax.bar(labels,means,color=colors,width=.65,yerr=yerr,capsize=5)
ax.set_ylabel('Gói TCP hợp pháp nhận được')
ax.set_title('Hiệu quả hai chế độ ứng phó – 8 seed cố định')
ax.grid(axis='y',alpha=.25)
ax.bar_label(bars,labels=['3.426','316,25','2.823,875','2.855,625'],padding=5,fontsize=9,fontweight='bold')
ax.text(.5,-.18,'Thanh sai số: khoảng 95% mô tả; không đại diện bất định tổng quát',transform=ax.transAxes,ha='center',fontsize=8.5)
fig.tight_layout(); fig.savefig(work/'figure_3_7_new.png',bbox_inches='tight'); plt.close(fig)

doc=Document(src)

# Captions/heading and scenario description.
replace_exact('3.5.2 Mô hình mạng và ba kịch bản','3.5.2 Mô hình mạng và bốn kịch bản đối chứng')
replace_starts('Để đánh giá có đối chứng, luận văn thiết lập ba kịch bản.',
'''Để đánh giá có đối chứng, thí nghiệm gồm bốn ô kịch bản: S0 là trạng thái nền không tấn công; S1 có tấn công nhưng không phòng vệ; S2-rate-limit kích hoạt giới hạn tốc độ; và S2-isolation cách ly nguồn tấn công. Hai biến thể S2 dùng cùng tải, seed và cấu hình mạng như S1 để bảo đảm so sánh theo cặp. Cấu trúc bốn ô kịch bản được tóm tắt trong Bảng 3.2.''')
for p in doc.paragraphs:
    if 'Bảng 3.2. Ba kịch bản mô phỏng đối chứng.' in p.text:
        setp(p,p.text.replace('Bảng 3.2. Ba kịch bản mô phỏng đối chứng.','Bảng 3.2. Bốn kịch bản mô phỏng đối chứng.'))

replace_starts('Công cụ mô phỏng là NS2 phiên bản 2.35',
'''Công cụ mô phỏng là NS-2 phiên bản 2.35 kết hợp mô-đun nOBS [23] trên môi trường Linux, sử dụng Simulator trace-all để lưu dấu vết gói và burst. Ma trận chính gồm 32 lượt chạy: bốn kịch bản nhân với tám seed cố định, mỗi lượt đo trong 5 giây và có 0,25 giây để tiêu thoát gói đang truyền. Backbone quang có tốc độ 400 Mb/s; tám nguồn UDP/CBR tấn công ở mức danh nghĩa 40 Mb/s/nguồn. Với bộ sinh số ngẫu nhiên của bản dựng này, hệ số tải của seed 1–8 co cụm quanh 0,8001–0,8011, tương ứng tải hiệu dụng khoảng 32,005–32,043 Mb/s/nguồn. Vì vậy, các khoảng 95% trong phần này chỉ mô tả biến thiên của tám seed cố định, không đại diện cho bất định tổng quát.''')

replace_starts('Kết quả so sánh giữa kịch bản nền và kịch bản tấn công cho thấy',
'''Kết quả chạy mới cho thấy tấn công làm số gói TCP hợp pháp nhận được giảm từ trung bình 3.426 gói ở S0 xuống 316,25 gói ở S1, tương ứng giảm 90,77%. Đồng thời, số burst được đưa vào mạng tăng từ 2.304 lên 8.056,125 (+249,66%) và tỷ lệ burst bị rớt tăng từ 0% lên 1,199%. Quan sát này phù hợp với cơ chế cạnh tranh tài nguyên trong nOBS: tải tấn công làm tăng mạnh số burst và xác suất thất bại đặt trước, khiến TCP hợp pháp suy giảm. Bảng 3.3 trình bày các chỉ số được tái suy trực tiếp từ raw trace của 32 lượt chạy; Hình 3.4 minh họa riêng đối chứng S0–S1.''')

# Response section: accurately describe simulated mechanism.
replace_starts('Cơ chế ứng phó thiết kế ở phần trên được hiện thực trong nOBS dưới dạng',
'''Hai chế độ ứng phó được hiện thực trong kịch bản NS-2/nOBS tại nút biên. Chế độ giới hạn tốc độ gắn bộ điều tiết token bucket với CIR 4 Mb/s cho từng nguồn tấn công; chế độ cách ly chuyển hướng gói của nguồn tấn công sang bộ nhận loại bỏ cục bộ. Trong ma trận này, hành động được kích hoạt sau độ trễ cấu hình 0,25 giây kể từ khi tấn công bắt đầu. Đây là mô phỏng cơ chế ứng phó sau phát hiện, không phải việc nhúng trực tiếp mô hình cây quyết định vào vòng chạy NS-2.''')
replace_starts('Một phép thử nội tại về cảnh báo sai đã được thực hiện',
'''Cấu hình giữ hai luồng TCP hợp pháp ở mức truy nhập 3 Mb/s và đặt CIR của nguồn bị giới hạn ở 4 Mb/s. Tám nguồn tấn công có tải danh nghĩa 40 Mb/s/nguồn, tương ứng khoảng 32,005–32,043 Mb/s/nguồn trong tám seed đã chọn. Việc tách định danh nguồn hợp pháp và nguồn tấn công là giả định của kịch bản ứng phó có oracle; do đó kết quả đo hiệu quả hành động giảm tải sau phát hiện, không tự thân chứng minh độ chính xác của bộ phát hiện hay tỷ lệ cảnh báo sai.''')
replace_starts('Bảng 3.6 trình bày thông lượng hợp pháp theo từng kịch bản',
'''Bảng 3.6 trình bày số gói TCP hợp pháp nhận được theo bốn ô kịch bản. Cả 32 lượt chạy đều kết thúc thành công; mỗi ô có tám seed. Toàn bộ raw trace được tái phân tích độc lập và đối chiếu khớp với metrics.json trước khi tổng hợp. Hình 3.7 minh họa trung bình và khoảng 95% mô tả trên tập seed cố định.''')
replace_starts('Kết quả mới cho thấy chế độ giới hạn tốc độ phục hồi được khoảng',
'''Giới hạn tốc độ nâng số gói TCP hợp pháp từ 316,25 lên 2.823,875 gói, đạt 82,42% mức nền và khôi phục 80,64% phần thông lượng bị mất giữa S0 và S1. Cách ly đạt 2.855,625 gói, bằng 83,35% mức nền và khôi phục 81,66% phần bị mất. So với S1, số burst offered giảm 64,10% ở chế độ giới hạn tốc độ và 70,91% ở chế độ cách ly; tỷ lệ burst drop tương ứng còn 0,043% và 0,053%.''')
replace_starts('Luận văn diễn giải kết quả này một cách thận trọng',
'''Hai chế độ đều phục hồi mạnh nhưng chưa trở lại hoàn toàn mức nền. Cách ly nhỉnh hơn giới hạn tốc độ về trung bình, song khoảng mô tả rộng hơn do seed 1 cho kết quả khác các seed còn lại. Vì seed 1–8 tạo dải hệ số tải rất hẹp, không nên suy rộng chênh lệch nhỏ giữa hai chế độ thành ưu thế tổng quát. Việc lựa chọn hành động vẫn là đánh đổi: cách ly cắt tải mạnh hơn nhưng rủi ro cao hơn nếu phát hiện sai; giới hạn tốc độ mềm hơn và duy trì một phần lưu lượng của nguồn bị nghi ngờ.''')
replace_starts('Tổng hợp lại, kết quả này khép lại nửa phần ứng phó',
'''Tổng hợp lại, ma trận 32 lượt chạy chứng minh trong cấu hình bảy nút đã khai báo rằng cả giới hạn tốc độ và cách ly đều giảm tải burst tấn công và khôi phục khoảng 80,64–81,66% phần thông lượng TCP hợp pháp bị mất. Kết luận này chỉ áp dụng cho cấu hình, tải và tám seed đã công bố; cần mở rộng seed, tô-pô và mẫu lưu lượng trước khi khái quát hóa.''')

replace_starts('Nhận định thứ hai là về bản chất đóng góp của luận văn.',
'''Nhận định thứ hai là về bản chất đóng góp của luận văn. Đóng góp không phải một thuật toán học máy mới, mà gồm: phát hiện và định lượng rò rỉ trong benchmark UCI; xây dựng benchmark NS-2 mức mạng không suy biến; và mô phỏng cơ chế ứng phó tại nút biên. Trong ma trận chạy mới, hai chế độ ứng phó khôi phục khoảng 80,64–81,66% phần thông lượng hợp pháp bị mất, tương đương 82,42–83,35% mức nền.''')
replace_starts('Nhận định thứ ba là phần trình bày trung thực các hạn chế.',
'''Nhận định thứ ba là phạm vi suy luận. Tấn công được mô hình hóa bằng UDP/CBR nhằm xấp xỉ hiệu ứng chiếm tài nguyên, chưa tái tạo giả mạo BHP ở mức giao thức. Hành động ứng phó dùng định danh nguồn đã biết và độ trễ kích hoạt cố định 0,25 giây, chưa tích hợp bộ phân loại trực tuyến vào NS-2. Ngoài ra, seed 1–8 tạo tải hiệu dụng rất hẹp; vì vậy khoảng 95% chỉ có ý nghĩa mô tả tập chạy, và kết quả cần được kiểm chứng trên nhiều seed, tô-pô và mẫu lưu lượng hơn.''')

replace_starts('Thứ nhất, luận văn đã định lượng được tác động của tấn công ngập lụt BHP lên mạng OBS',
'''Thứ nhất, luận văn định lượng tác động của kịch bản tấn công xấp xỉ ngập lụt BHP bằng NS-2.35/nOBS. Trên tám seed cố định, số gói TCP hợp pháp giảm từ 3.426 xuống 316,25 gói (−90,77%), burst offered tăng 249,66% và tỷ lệ burst drop tăng từ 0% lên 1,199%. Kết quả cho thấy cạnh tranh tài nguyên đặt trước làm suy giảm nghiêm trọng lưu lượng hợp pháp trong cấu hình khảo sát (mục tiêu 1).''')
replace_starts('Thứ năm, luận văn mô phỏng toàn bộ cơ chế đề xuất bằng NS2 phiên bản',
'''Thứ năm, luận văn chạy ma trận 32 lượt bằng NS-2.35/nOBS trên bốn ô kịch bản và tám seed cố định. Cả 32 lượt đều thành công và raw trace được tái phân tích độc lập. Giới hạn tốc độ đạt 82,42% mức nền, khôi phục 80,64% phần thông lượng bị mất; cách ly đạt 83,35% mức nền, khôi phục 81,66%. Kết quả chứng minh hiệu quả của hành động giảm tải sau phát hiện trong cấu hình nghiên cứu, nhưng không đồng nghĩa bộ phát hiện đã được nhúng trực tiếp vào NS-2 (mục tiêu 5).''')

# Conclusion paragraph if present.
replace_starts('Về bài toán ứng phó, luận văn chứng minh bằng định lượng',
'''Về bài toán ứng phó, ma trận 32 lượt chạy NS-2.35/nOBS cho thấy giới hạn tốc độ và cách ly đều phục hồi mạnh lưu lượng TCP hợp pháp sau tấn công, lần lượt đạt 82,42% và 83,35% mức nền. Hai chế độ khôi phục 80,64% và 81,66% phần thông lượng bị mất. Đây là kết quả mô phỏng cho cấu hình bảy nút và tám seed cố định; hành động được kích hoạt sau độ trễ cấu hình 0,25 giây với định danh nguồn đã biết.''')

# Tables.
# Table 3: four scenario rows
T=doc.tables[3]
rows=[
['Kịch bản','Lưu lượng hợp pháp','Tấn công','Phòng vệ'],
['Nền (S0)','Có','Không','Không'],
['Tấn công (S1)','Có','Có','Không'],
['Ứng phó (S2-rate-limit)','Có','Có','Giới hạn tốc độ, CIR 4 Mb/s'],
['Ứng phó (S2-isolation)','Có','Có','Cách ly tại nút biên'],
]
while len(T.rows)<len(rows): T.add_row()
for i,row in enumerate(rows):
    for j,val in enumerate(row): T.cell(i,j).text=val

T=doc.tables[4]
rows=[
['Chỉ số','S0 – nền','S1 – tấn công','Thay đổi S1/S0','Ghi chú'],
['Gói TCP hợp pháp','3.426','316,25','−90,77%','Trung bình 8 seed'],
['Byte TCP hợp pháp','3.561.040','326.900','−90,82%','Trung bình 8 seed'],
['Burst offered','2.304','8.056,125','+249,66%','Tái suy từ raw trace'],
['Tỷ lệ burst drop','0%','1,199%','+1,199 điểm %','Drop tường minh trong trace'],
]
for i,row in enumerate(rows):
    for j,val in enumerate(row): T.cell(i,j).text=val

T=doc.tables[7]
rows=[
['Kịch bản','n','Gói TCP hợp pháp (TB)','Khoảng 95% mô tả','So với nền'],
['Nền (S0)','8','3.426','[3.426; 3.426]','100%'],
['Tấn công (S1)','8','316,25','[278,138; 354,362]','9,23%'],
['Giới hạn tốc độ (S2)','8','2.823,875','[2.814,507; 2.833,243]','82,42%'],
['Cách ly (S2)','8','2.855,625','[2.686,259; 3.024,991]','83,35%'],
]
for i,row in enumerate(rows):
    for j,val in enumerate(row): T.cell(i,j).text=val

# Save DOCX first.
doc.save(out)

# Replace image blobs while preserving drawing layout/relationships.
def replace_zip_members(docx_path, replacements):
    tmp=docx_path.with_suffix('.tmp.docx')
    with zipfile.ZipFile(docx_path,'r') as zin, zipfile.ZipFile(tmp,'w',zipfile.ZIP_DEFLATED) as zout:
        names=set(zin.namelist())
        for item in zin.infolist():
            data=zin.read(item.filename)
            if item.filename in replacements:
                data=Path(replacements[item.filename]).read_bytes()
            zout.writestr(item,data)
        missing=set(replacements)-names
        if missing: raise RuntimeError(f'missing media members: {missing}')
    os.replace(tmp,docx_path)
replace_zip_members(out,{
    'word/media/image26.png': work/'figure_3_4_new.png',
    'word/media/image29.png': work/'figure_3_7_new.png',
})
print(out)
