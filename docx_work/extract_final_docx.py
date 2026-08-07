from pathlib import Path
from docx import Document

src = Path("LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx")
out = Path("docx_work/final_extracted.txt")
doc = Document(src)
with out.open("w", encoding="utf-8") as handle:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            handle.write(f"P{index}\t{paragraph.text}\n")
    for table_index, table in enumerate(doc.tables):
        for row_index, row in enumerate(table.rows):
            handle.write(
                f"T{table_index}R{row_index}\t"
                + " | ".join(cell.text for cell in row.cells)
                + "\n"
            )
print(f"{out} paragraphs={len(doc.paragraphs)} tables={len(doc.tables)}")
