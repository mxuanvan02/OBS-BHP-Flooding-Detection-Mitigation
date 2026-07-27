from pathlib import Path
from zipfile import ZipFile
from docx import Document
import sys

p = Path(sys.argv[1]).resolve()
with ZipFile(p) as z:
    bad = z.testzip()
    assert bad is None, f"corrupt zip member: {bad}"
    assert "word/document.xml" in z.namelist()

doc = Document(p)
parts = [paragraph.text for paragraph in doc.paragraphs]
for table in doc.tables:
    for row in table.rows:
        parts.extend(cell.text for cell in row.cells)
text = "\n".join(parts)

required = [
    "48.678",
    "24.307,5",
    "50.623.120",
    "25.277.800",
    "50,07%",
    "1000 Mb/s",
    "12 Mb/s/nguồn",
    "control-only BHP",
    "100% mức S0",
]
for needle in required:
    assert needle in text, f"missing required updated content: {needle}"

forbidden = [
    "82.568",
    "38.281",
    "53.078",
    "84.834",
    "53,6 phần trăm",
    "DET_DELAY = 0,10",
    "tốc độ khoảng mười hai megabit",
    "105,1 phần trăm",
    "102,7 phần trăm",
    "3.426",
    "316,25",
    "2.823,875",
    "2.855,625",
    "90,77%",
    "400 Mb/s",
    "40 Mb/s/nguồn",
    "guard tắt",
]
found = [needle for needle in forbidden if needle in text]
assert not found, f"obsolete content remains: {found}"

print(
    f"DOCX_STRUCTURE_OK paragraphs={len(doc.paragraphs)} "
    f"tables={len(doc.tables)} bytes={p.stat().st_size}"
)
print("UPDATED_CONTENT_OK")
