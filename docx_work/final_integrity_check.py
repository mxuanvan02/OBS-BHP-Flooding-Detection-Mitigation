#!/usr/bin/env python3
from pathlib import Path
from zipfile import ZipFile
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import hashlib

BASE = Path(__file__).resolve().parents[1]
FILES = [
    BASE / "LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx",
    BASE / "LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726_HIGHLIGHT.docx",
]

results = []
for path in FILES:
    assert path.is_file(), f"missing file: {path}"
    data = path.read_bytes()
    with ZipFile(path) as archive:
        bad_member = archive.testzip()
        names = archive.namelist()
        media = [name for name in names if name.startswith("word/media/")]
        assert bad_member is None, f"corrupt ZIP member in {path}: {bad_member}"
        assert "word/document.xml" in names, f"missing word/document.xml: {path}"

    document = Document(path)
    paragraph_highlights = sum(
        1
        for paragraph in document.paragraphs
        for run in paragraph.runs
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW
    )
    table_highlights = sum(
        1
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW
    )
    result = {
        "path": str(path),
        "bytes": len(data),
        "sha256": hashlib.sha256(data).hexdigest(),
        "zip_bad": bad_member,
        "media": len(media),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "paragraph_highlights": paragraph_highlights,
        "table_highlights": table_highlights,
    }
    results.append(result)
    print(" ".join(f"{key}={value}" for key, value in result.items()))

normal, highlighted = results
assert normal["paragraph_highlights"] == 0, "normal DOCX unexpectedly contains yellow paragraph highlights"
assert highlighted["paragraph_highlights"] + highlighted["table_highlights"] > 0, "highlight DOCX has no yellow changes"
print("INTEGRITY_AND_HIGHLIGHT_OK")
