from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

base = Path(__file__).resolve().parents[1]
old_path = base / 'deliverables/LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
new_path = base / 'LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726.docx'
out_path = base / 'LuanVan_ThS_NguyenQuangTin_CAPNHAT_KETQUA_NS2_20260726_HIGHLIGHT.docx'

old = Document(str(old_path))
new = Document(str(new_path))
old_paras = [p.text for p in old.paragraphs]
old_tables = [[[c.text for c in row.cells] for row in t.rows] for t in old.tables]

changed_paras = 0
changed_cells = 0

def highlight_paragraph(p):
    for run in p.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    # Paragraphs containing only drawings have no runs; highlight caption/marker text instead.
    if not p.runs and p.text:
        for run in p.runs:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

def shade_cell(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), 'FFF2CC')

for i, p in enumerate(new.paragraphs):
    old_text = old_paras[i] if i < len(old_paras) else None
    if p.text and p.text != old_text:
        highlight_paragraph(p)
        changed_paras += 1

for ti, table in enumerate(new.tables):
    old_table = old_tables[ti] if ti < len(old_tables) else []
    for ri, row in enumerate(table.rows):
        old_row = old_table[ri] if ri < len(old_table) else []
        for ci, cell in enumerate(row.cells):
            old_text = old_row[ci] if ci < len(old_row) else None
            if cell.text and cell.text != old_text:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                shade_cell(cell)
                changed_cells += 1

# Make the review legend explicit, without changing the scientific content.
first = new.paragraphs[0]
legend = first.insert_paragraph_before('Ghi chú rà soát: phần được cập nhật trong lần chạy NS-2.35/nOBS mới được đánh dấu nền vàng.')
for run in legend.runs:
    run.bold = True
    run.font.highlight_color = WD_COLOR_INDEX.YELLOW

new.save(str(out_path))
print(f'created={out_path}')
print(f'highlighted_paragraphs={changed_paras} highlighted_cells={changed_cells}')
